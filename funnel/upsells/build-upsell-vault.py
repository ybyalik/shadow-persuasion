#!/usr/bin/env python3
"""
The Shadow Persuasion Vault — 250 Field-Tested Techniques.
Upsell #1 ($37). 10 parts x 25 techniques, each with breakdown.

Cover layout matches the 4 bonus covers. Body is a 6x9" docx
with header logo + footer page numbers, chapter breaks per part.
"""

from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Load all 250 techniques from the 3 data files
import sys
HERE = Path(__file__).resolve().parent
sys.path.insert(0, str(HERE))
from vault_techniques import PARTS as PARTS_1       # Part 1
from vault_parts_2_5 import PARTS_2_5                # Parts 2-5
from vault_parts_6_10 import PARTS_6_10              # Parts 6-10

ALL_PARTS = PARTS_1 + PARTS_2_5 + PARTS_6_10

REPO = HERE.parents[1]
LOGO = REPO / "public" / "logo.png"
FONT_TTF = REPO / "funnel" / "assets" / "SpecialElite-Regular.ttf"
COVER_OUT = REPO / "funnel" / "assets" / "cover-upsell-vault.png"
DOCX_OUT = REPO / "funnel" / "exports" / "upsell-shadow-persuasion-vault.docx"

BODY_FONT = "Georgia"
TITLE_FONT = "Georgia"

# Brand colors
CREAM = (0xF4, 0xEC, 0xD8)
DARK = (0x1A, 0x1A, 0x1A)
BROWN = (0x5C, 0x3A, 0x1E)
GOLD = (0xD4, 0xA0, 0x17)

DARK_RGB = RGBColor(0x1A, 0x1A, 0x1A)
BROWN_RGB = RGBColor(0x5C, 0x3A, 0x1E)
GOLD_RGB = RGBColor(0xD4, 0xA0, 0x17)
GRAY_RGB = RGBColor(0x6B, 0x5B, 0x3E)


# ========== COVER ==========

def build_cover():
    WIDTH, HEIGHT = 1800, 2700
    img = Image.new("RGB", (WIDTH, HEIGHT), CREAM)
    draw = ImageDraw.Draw(img)

    label_font = ImageFont.truetype(str(FONT_TTF), 40)
    title_font = ImageFont.truetype(str(FONT_TTF), 110)
    subtitle_font = ImageFont.truetype(str(FONT_TTF), 58)
    badge_font = ImageFont.truetype(str(FONT_TTF), 48)
    author_font = ImageFont.truetype(str(FONT_TTF), 52)

    # Logo
    logo = Image.open(str(LOGO)).convert("RGBA")
    scale = 900 / logo.width
    logo_resized = logo.resize((int(logo.width * scale), int(logo.height * scale)), Image.LANCZOS)
    logo_x = (WIDTH - logo_resized.width) // 2
    logo_y = 380
    img.paste(logo_resized, (logo_x, logo_y), logo_resized)

    # Exclusive label
    label_text = "// THE COMPLETE VAULT //"
    bbox = draw.textbbox((0, 0), label_text, font=label_font)
    label_y = logo_y + logo_resized.height + 80
    draw.text(((WIDTH - (bbox[2] - bbox[0])) // 2, label_y), label_text, fill=BROWN, font=label_font)

    # Title
    title_y = label_y + 120
    title_line_spacing = 135
    title_lines = ["SHADOW", "PERSUASION", "VAULT"]
    for i, line in enumerate(title_lines):
        bbox = draw.textbbox((0, 0), line, font=title_font)
        draw.text(((WIDTH - (bbox[2] - bbox[0])) // 2, title_y + i * title_line_spacing),
                  line, fill=DARK, font=title_font)

    # Gold rule
    rule_y = title_y + (len(title_lines) * title_line_spacing) + 60
    draw.rectangle([(WIDTH // 2 - 200, rule_y), (WIDTH // 2 + 200, rule_y + 3)], fill=GOLD)

    # Subtitle
    subtitle_lines = [
        "250 Field-Tested Techniques",
        "With Full Deployment Breakdowns",
    ]
    sub_y = rule_y + 80
    for line in subtitle_lines:
        bbox = draw.textbbox((0, 0), line, font=subtitle_font)
        draw.text(((WIDTH - (bbox[2] - bbox[0])) // 2, sub_y), line, fill=BROWN, font=subtitle_font)
        sub_y += 90

    # Big 250 badge — black box with gold number
    badge_w, badge_h = 360, 140
    badge_x = (WIDTH - badge_w) // 2
    badge_y = sub_y + 90
    draw.rectangle([(badge_x, badge_y), (badge_x + badge_w, badge_y + badge_h)], fill=DARK)
    big_num_font = ImageFont.truetype(str(FONT_TTF), 90)
    num_text = "250"
    bbox = draw.textbbox((0, 0), num_text, font=big_num_font)
    nx = badge_x + (badge_w - (bbox[2] - bbox[0])) // 2
    ny = badge_y + (badge_h - (bbox[3] - bbox[1])) // 2 - 15
    draw.text((nx, ny), num_text, fill=GOLD, font=big_num_font)

    tag_font = ImageFont.truetype(str(FONT_TTF), 34)
    tag_text = "TECHNIQUES"
    bbox = draw.textbbox((0, 0), tag_text, font=tag_font)
    draw.text(((WIDTH - (bbox[2] - bbox[0])) // 2, badge_y + badge_h + 30), tag_text, fill=BROWN, font=tag_font)

    # Author — bottom
    author_text = "NATE HARLAN"
    bbox = draw.textbbox((0, 0), author_text, font=author_font)
    draw.text(((WIDTH - (bbox[2] - bbox[0])) // 2, HEIGHT - 280), author_text, fill=DARK, font=author_font)

    # Border
    draw.rectangle([(60, 60), (WIDTH - 60, HEIGHT - 60)], outline=BROWN, width=2)

    COVER_OUT.parent.mkdir(parents=True, exist_ok=True)
    img.save(str(COVER_OUT), "PNG", dpi=(300, 300))
    print(f"Cover saved: {COVER_OUT}")


# ========== DOCX HELPERS ==========

def add_page_number(paragraph):
    run = paragraph.add_run()
    el = OxmlElement("w:fldChar")
    el.set(qn("w:fldCharType"), "begin")
    run._r.append(el)
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = "PAGE"
    run._r.append(it)
    el = OxmlElement("w:fldChar")
    el.set(qn("w:fldCharType"), "end")
    run._r.append(el)
    run.font.name = BODY_FONT
    run.font.size = Pt(10)
    run.font.color.rgb = BROWN_RGB


def set_page_break_before(p):
    pPr = p._p.get_or_add_pPr()
    pPr.append(OxmlElement("w:pageBreakBefore"))


# ========== DOCX BUILD ==========

def build_docx():
    doc = Document()

    # ==== SECTION 1: COVER (zero-margin full bleed) ====
    cover_section = doc.sections[0]
    cover_section.page_width = Inches(6)
    cover_section.page_height = Inches(9)
    cover_section.top_margin = Inches(0)
    cover_section.bottom_margin = Inches(0)
    cover_section.left_margin = Inches(0)
    cover_section.right_margin = Inches(0)
    cover_section.header_distance = Inches(0)
    cover_section.footer_distance = Inches(0)

    cover_p = doc.add_paragraph()
    cover_p.paragraph_format.space_before = Pt(0)
    cover_p.paragraph_format.space_after = Pt(0)
    cover_p.paragraph_format.line_spacing = 1.0
    run = cover_p.add_run()
    run.add_picture(str(COVER_OUT), width=Inches(6), height=Inches(9))

    # ==== SECTION 2: BODY ====
    body = doc.add_section(WD_SECTION_START.NEW_PAGE)
    body.page_width = Inches(6)
    body.page_height = Inches(9)
    body.top_margin = Inches(1.0)
    body.bottom_margin = Inches(0.85)
    body.left_margin = Inches(0.75)
    body.right_margin = Inches(0.75)
    body.header_distance = Inches(0.3)
    body.footer_distance = Inches(0.4)
    body.header.is_linked_to_previous = False
    body.footer.is_linked_to_previous = False

    # Header: logo right-aligned
    h = body.header.paragraphs[0]
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hrun = h.add_run()
    hrun.add_picture(str(LOGO), width=Inches(0.9))

    # Footer: page number centered
    f = body.footer.paragraphs[0]
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(f)

    # ==== TITLE PAGE ====
    # Big title block on the first body page
    for spacer in range(3):
        doc.add_paragraph()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_after = Pt(6)
    r = t.add_run("THE SHADOW PERSUASION VAULT")
    r.font.name = TITLE_FONT
    r.font.size = Pt(24)
    r.bold = True
    r.font.color.rgb = DARK_RGB

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st.paragraph_format.space_after = Pt(20)
    r = st.add_run("250 Field-Tested Techniques With Full Deployment Breakdowns")
    r.font.name = TITLE_FONT
    r.font.size = Pt(13)
    r.italic = True
    r.font.color.rgb = BROWN_RGB

    by = doc.add_paragraph()
    by.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = by.add_run("by Nate Harlan")
    r.font.name = TITLE_FONT
    r.font.size = Pt(12)
    r.font.color.rgb = GRAY_RGB

    # ==== HOW TO USE ====
    h = doc.add_paragraph()
    set_page_break_before(h)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.paragraph_format.space_before = Pt(24)
    h.paragraph_format.space_after = Pt(16)
    r = h.add_run("HOW TO USE THIS VAULT")
    r.font.name = TITLE_FONT
    r.font.size = Pt(20)
    r.bold = True
    r.font.color.rgb = DARK_RGB

    intro_paragraphs = [
        "This is not a book to read cover-to-cover. It is a reference library. Each of the 250 techniques stands on its own and can be deployed in isolation. You will use a small subset of these in any given month. Over a year you will use most of them. Over a career you will use all of them.",
        "The techniques are organized into ten parts that mirror the four-part system in Shadow Persuasion: Openers and Detector Disablers. Rapport Without Rapport. Making Them Persuade Themselves. Silence, Pauses, and Timing. Reframing and Frame Control. Handling Objections. Defensive and Anti-Manipulation. Reading the Room. Closes Without Closing. Post-Conversation Lock-Ins.",
        "Each technique has the same structure. USE WHEN tells you the exact situation the move applies to. HOW IT WORKS explains the mechanism. DEPLOY gives you a concrete example you can copy. COUNTER tells you what to do when someone uses the technique on you.",
        "Read one part at a time. Practice two or three techniques from each part in low-stakes conversations this week. Then rotate. Within a month the ones you use most will be second-nature. Within six months, you will be using moves you read here without remembering that you read them here.",
        "One warning. Most of these techniques work by disabling the persuasion detector in the other person's head. That means they will not consciously notice you are using them. That is a feature, not a bug. It also means you are responsible for how you use them. These are tools. Tools are neutral. People are not.",
    ]
    for para in intro_paragraphs:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.35
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.first_line_indent = Inches(0.25)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(para)
        r.font.name = BODY_FONT
        r.font.size = Pt(11)
        r.font.color.rgb = DARK_RGB

    # ==== TABLE OF CONTENTS ====
    toc_h = doc.add_paragraph()
    set_page_break_before(toc_h)
    toc_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_h.paragraph_format.space_before = Pt(24)
    toc_h.paragraph_format.space_after = Pt(16)
    r = toc_h.add_run("CONTENTS")
    r.font.name = TITLE_FONT
    r.font.size = Pt(20)
    r.bold = True
    r.font.color.rgb = DARK_RGB

    running_num = 1
    for part in ALL_PARTS:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"PART {part['num']:02d}  \u2014  {part['title']}")
        r.font.name = TITLE_FONT
        r.font.size = Pt(11)
        r.bold = True
        r.font.color.rgb = DARK_RGB

        sub = doc.add_paragraph()
        sub.paragraph_format.left_indent = Inches(0.2)
        sub.paragraph_format.space_after = Pt(8)
        r = sub.add_run(f"Techniques {running_num} to {running_num + len(part['techniques']) - 1}  \u2014  {part['subtitle']}")
        r.font.name = TITLE_FONT
        r.font.size = Pt(9.5)
        r.italic = True
        r.font.color.rgb = GRAY_RGB
        running_num += len(part['techniques'])

    # ==== PARTS + TECHNIQUES ====
    technique_num = 1
    for part in ALL_PARTS:
        # Part divider page
        pd = doc.add_paragraph()
        set_page_break_before(pd)
        pd.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pd.paragraph_format.space_before = Pt(120)
        pd.paragraph_format.space_after = Pt(8)
        r = pd.add_run(f"PART {part['num']:02d}")
        r.font.name = TITLE_FONT
        r.font.size = Pt(16)
        r.font.color.rgb = GOLD_RGB

        pt = doc.add_paragraph()
        pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pt.paragraph_format.space_after = Pt(10)
        r = pt.add_run(part['title'])
        r.font.name = TITLE_FONT
        r.font.size = Pt(26)
        r.bold = True
        r.font.color.rgb = DARK_RGB

        ps = doc.add_paragraph()
        ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ps.paragraph_format.space_after = Pt(16)
        r = ps.add_run(part['subtitle'])
        r.font.name = TITLE_FONT
        r.font.size = Pt(11)
        r.italic = True
        r.font.color.rgb = BROWN_RGB

        # Techniques in this part
        for tech in part['techniques']:
            # Technique header (don't page-break before each — flow them)
            th = doc.add_paragraph()
            th.paragraph_format.space_before = Pt(14)
            th.paragraph_format.space_after = Pt(4)
            th.paragraph_format.keep_with_next = True
            r = th.add_run(f"#{technique_num:03d}  \u2014  {tech['name']}")
            r.font.name = TITLE_FONT
            r.font.size = Pt(14)
            r.bold = True
            r.font.color.rgb = DARK_RGB

            # Metadata line
            meta = doc.add_paragraph()
            meta.paragraph_format.space_after = Pt(6)
            meta.paragraph_format.keep_with_next = True
            r = meta.add_run(f"CATEGORY: {tech['category']}   \u00b7   RISK: {tech['risk']}   \u00b7   DOMAINS: {tech['domains']}")
            r.font.name = BODY_FONT
            r.font.size = Pt(8.5)
            r.font.color.rgb = GRAY_RGB

            # Use when
            for label, content_key in [("USE WHEN", "use_when"),
                                         ("HOW IT WORKS", "how"),
                                         ("DEPLOY", "deploy"),
                                         ("COUNTER", "counter")]:
                if content_key in tech:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(4)
                    p.paragraph_format.left_indent = Inches(0.0)
                    p.paragraph_format.line_spacing = 1.3
                    r1 = p.add_run(f"{label}: ")
                    r1.font.name = BODY_FONT
                    r1.font.size = Pt(9)
                    r1.bold = True
                    r1.font.color.rgb = GOLD_RGB
                    r2 = p.add_run(tech[content_key])
                    r2.font.name = BODY_FONT
                    r2.font.size = Pt(10)
                    r2.font.color.rgb = DARK_RGB

            technique_num += 1

    # ==== CLOSING PAGE ====
    c = doc.add_paragraph()
    set_page_break_before(c)
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_before = Pt(180)
    c.paragraph_format.space_after = Pt(16)
    r = c.add_run("THAT'S THE VAULT.")
    r.font.name = TITLE_FONT
    r.font.size = Pt(22)
    r.bold = True
    r.font.color.rgb = DARK_RGB

    c2 = doc.add_paragraph()
    c2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c2.paragraph_format.space_after = Pt(14)
    r = c2.add_run("250 techniques. 10 categories. One purpose:")
    r.font.name = TITLE_FONT
    r.font.size = Pt(11)
    r.font.color.rgb = BROWN_RGB

    c3 = doc.add_paragraph()
    c3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c3.paragraph_format.space_after = Pt(20)
    r = c3.add_run("Win the conversations that actually decide your life,\nwithout the other person ever realizing you tried to.")
    r.font.name = TITLE_FONT
    r.font.size = Pt(12)
    r.italic = True
    r.font.color.rgb = DARK_RGB

    c4 = doc.add_paragraph()
    c4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c4.add_run("\u2014  Nate Harlan")
    r.font.name = TITLE_FONT
    r.font.size = Pt(11)
    r.font.color.rgb = GRAY_RGB

    DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(DOCX_OUT))
    print(f"Docx saved: {DOCX_OUT}")
    print(f"Size: {DOCX_OUT.stat().st_size} bytes")
    print(f"Total techniques: {technique_num - 1}")


if __name__ == "__main__":
    build_cover()
    build_docx()
