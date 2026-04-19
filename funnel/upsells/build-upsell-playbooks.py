#!/usr/bin/env python3
"""
Upsell #1: The Shadow Persuasion Situation Playbooks ($47).

20 situation-specific playbooks, 2-3 pages each, grouped into 5 parts:
Career, Dating, Long-Term Relationships, Endings & Recovery, Family & Life.
"""

from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import sys
HERE = Path(__file__).resolve().parent
sys.path.insert(0, str(HERE))
from playbooks_data import PLAYBOOKS

REPO = HERE.parents[1]
LOGO = REPO / "public" / "logo.png"
FONT_TTF = REPO / "funnel" / "assets" / "SpecialElite-Regular.ttf"
COVER_OUT = REPO / "funnel" / "assets" / "cover-upsell-playbooks.png"
DOCX_OUT = REPO / "funnel" / "exports" / "upsell-situation-playbooks.docx"

BODY_FONT = "Georgia"
TITLE_FONT = "Georgia"

CREAM = (0xF4, 0xEC, 0xD8)
DARK = (0x1A, 0x1A, 0x1A)
BROWN = (0x5C, 0x3A, 0x1E)
GOLD = (0xD4, 0xA0, 0x17)

DARK_RGB = RGBColor(0x1A, 0x1A, 0x1A)
BROWN_RGB = RGBColor(0x5C, 0x3A, 0x1E)
GOLD_RGB = RGBColor(0xD4, 0xA0, 0x17)
GRAY_RGB = RGBColor(0x6B, 0x5B, 0x3E)


# ========== COVER ==========

def build_cover():
    WIDTH, HEIGHT = 1800, 2700
    img = Image.new("RGB", (WIDTH, HEIGHT), CREAM)
    draw = ImageDraw.Draw(img)

    label_font = ImageFont.truetype(str(FONT_TTF), 40)
    title_font = ImageFont.truetype(str(FONT_TTF), 110)
    subtitle_font = ImageFont.truetype(str(FONT_TTF), 56)
    badge_font = ImageFont.truetype(str(FONT_TTF), 40)
    author_font = ImageFont.truetype(str(FONT_TTF), 52)

    # Logo
    logo = Image.open(str(LOGO)).convert("RGBA")
    scale = 900 / logo.width
    logo_resized = logo.resize((int(logo.width * scale), int(logo.height * scale)), Image.LANCZOS)
    logo_x = (WIDTH - logo_resized.width) // 2
    logo_y = 380
    img.paste(logo_resized, (logo_x, logo_y), logo_resized)

    # Label
    label_text = "// THE FIELD LIBRARY //"
    bbox = draw.textbbox((0, 0), label_text, font=label_font)
    label_y = logo_y + logo_resized.height + 80
    draw.text(((WIDTH - (bbox[2] - bbox[0])) // 2, label_y), label_text, fill=BROWN, font=label_font)

    # Title — 3 lines
    title_y = label_y + 120
    title_line_spacing = 135
    title_lines = ["THE SITUATION", "PLAYBOOKS"]
    for i, line in enumerate(title_lines):
        bbox = draw.textbbox((0, 0), line, font=title_font)
        draw.text(((WIDTH - (bbox[2] - bbox[0])) // 2, title_y + i * title_line_spacing),
                  line, fill=DARK, font=title_font)

    # Gold rule
    rule_y = title_y + (len(title_lines) * title_line_spacing) + 60
    draw.rectangle([(WIDTH // 2 - 200, rule_y), (WIDTH // 2 + 200, rule_y + 3)], fill=GOLD)

    # Subtitle
    subtitle_lines = [
        "20 Playbooks For The Conversations",
        "That Actually Decide Your Life",
    ]
    sub_y = rule_y + 80
    for line in subtitle_lines:
        bbox = draw.textbbox((0, 0), line, font=subtitle_font)
        draw.text(((WIDTH - (bbox[2] - bbox[0])) // 2, sub_y), line, fill=BROWN, font=subtitle_font)
        sub_y += 90

    # Big 20 badge — black box with gold number
    badge_w, badge_h = 360, 130
    badge_x = (WIDTH - badge_w) // 2
    badge_y = sub_y + 100
    draw.rectangle([(badge_x, badge_y), (badge_x + badge_w, badge_y + badge_h)], fill=DARK)
    big_num_font = ImageFont.truetype(str(FONT_TTF), 90)
    num_text = "20"
    bbox = draw.textbbox((0, 0), num_text, font=big_num_font)
    nx = badge_x + (badge_w - (bbox[2] - bbox[0])) // 2
    ny = badge_y + (badge_h - (bbox[3] - bbox[1])) // 2 - 15
    draw.text((nx, ny), num_text, fill=GOLD, font=big_num_font)

    tag_font = ImageFont.truetype(str(FONT_TTF), 32)
    tag_text = "PLAYBOOKS"
    bbox = draw.textbbox((0, 0), tag_text, font=tag_font)
    draw.text(((WIDTH - (bbox[2] - bbox[0])) // 2, badge_y + badge_h + 28), tag_text, fill=BROWN, font=tag_font)

    # Author
    author_text = "NATE HARLAN"
    bbox = draw.textbbox((0, 0), author_text, font=author_font)
    draw.text(((WIDTH - (bbox[2] - bbox[0])) // 2, HEIGHT - 280), author_text, fill=DARK, font=author_font)

    # Border
    draw.rectangle([(60, 60), (WIDTH - 60, HEIGHT - 60)], outline=BROWN, width=2)

    COVER_OUT.parent.mkdir(parents=True, exist_ok=True)
    img.save(str(COVER_OUT), "PNG", dpi=(300, 300))
    print(f"Cover saved: {COVER_OUT}")


# ========== DOCX HELPERS ==========

def add_page_number(paragraph):
    run = paragraph.add_run()
    el = OxmlElement("w:fldChar")
    el.set(qn("w:fldCharType"), "begin")
    run._r.append(el)
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = "PAGE"
    run._r.append(it)
    el = OxmlElement("w:fldChar")
    el.set(qn("w:fldCharType"), "end")
    run._r.append(el)
    run.font.name = BODY_FONT
    run.font.size = Pt(10)
    run.font.color.rgb = BROWN_RGB


def set_page_break_before(p):
    pPr = p._p.get_or_add_pPr()
    pPr.append(OxmlElement("w:pageBreakBefore"))


# ========== DOCX BUILD ==========

def build_docx():
    doc = Document()

    # ==== COVER SECTION ====
    s = doc.sections[0]
    s.page_width = Inches(6)
    s.page_height = Inches(9)
    s.top_margin = Inches(0)
    s.bottom_margin = Inches(0)
    s.left_margin = Inches(0)
    s.right_margin = Inches(0)
    s.header_distance = Inches(0)
    s.footer_distance = Inches(0)

    cover_p = doc.add_paragraph()
    cover_p.paragraph_format.space_before = Pt(0)
    cover_p.paragraph_format.space_after = Pt(0)
    cover_p.paragraph_format.line_spacing = 1.0
    run = cover_p.add_run()
    run.add_picture(str(COVER_OUT), width=Inches(6), height=Inches(9))

    # ==== BODY SECTION ====
    b = doc.add_section(WD_SECTION_START.NEW_PAGE)
    b.page_width = Inches(6)
    b.page_height = Inches(9)
    b.top_margin = Inches(1.0)
    b.bottom_margin = Inches(0.85)
    b.left_margin = Inches(0.75)
    b.right_margin = Inches(0.75)
    b.header_distance = Inches(0.3)
    b.footer_distance = Inches(0.4)
    b.header.is_linked_to_previous = False
    b.footer.is_linked_to_previous = False

    h = b.header.paragraphs[0]
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hrun = h.add_run()
    hrun.add_picture(str(LOGO), width=Inches(0.9))

    f = b.footer.paragraphs[0]
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(f)

    # ==== TITLE PAGE ====
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("THE SITUATION PLAYBOOKS")
    r.font.name = TITLE_FONT
    r.font.size = Pt(24)
    r.bold = True
    r.font.color.rgb = DARK_RGB

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    r = p.add_run("20 Playbooks For The Conversations That Actually Decide Your Life")
    r.font.name = TITLE_FONT
    r.font.size = Pt(12)
    r.italic = True
    r.font.color.rgb = BROWN_RGB

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("by Nate Harlan")
    r.font.name = TITLE_FONT
    r.font.size = Pt(11)
    r.font.color.rgb = GRAY_RGB

    # ==== HOW TO USE ====
    p = doc.add_paragraph()
    set_page_break_before(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("HOW TO USE THESE PLAYBOOKS")
    r.font.name = TITLE_FONT
    r.font.size = Pt(18)
    r.bold = True
    r.font.color.rgb = DARK_RGB

    intro = [
        "This is not a book to read in order. It is a library. When one of these twenty conversations is coming up in your life \u2014 and it will \u2014 you flip to that playbook, read it in about 10 minutes, fill out the Pre-Conversation Briefing, and walk into the room prepared at a level most people never reach.",
        "Each playbook is structured identically, so once you\u2019ve run one, the rest are fast. THE STAKES tells you what\u2019s actually at risk. THE DYNAMICS tells you what makes this specific conversation different from others. YOUR 3 TACTICS points you to the exact moves from Shadow Persuasion that work best here, with page references. THE OPENING LINE is word-for-word. MID-CONVERSATION MOVES are the 3 pivots that carry you through. THE 2 OBJECTIONS gives you pre-drafted responses to the pushback you\u2019ll actually get. YOUR EXIT MOVE is how you close. ONE MISTAKE TO AVOID names the specific trap. THE DEBRIEF QUESTION is what you ask yourself within two hours after.",
        "These playbooks work best paired with the Pre-Conversation Briefing. The playbook gives you the situation-specific framework. The briefing turns it into your specific conversation. Together they do more than either one alone.",
        "One warning. These playbooks cover twenty situations. Your life will contain conversations that don\u2019t exactly match any of them. That\u2019s not a flaw. The playbooks teach a shape of thinking that generalizes. After running three or four of them, you\u2019ll start writing your own for conversations we didn\u2019t include.",
    ]
    for text in intro:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.35
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.first_line_indent = Inches(0.25)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(text)
        r.font.name = BODY_FONT
        r.font.size = Pt(11)
        r.font.color.rgb = DARK_RGB

    # ==== CONTENTS ====
    p = doc.add_paragraph()
    set_page_break_before(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("CONTENTS")
    r.font.name = TITLE_FONT
    r.font.size = Pt(20)
    r.bold = True
    r.font.color.rgb = DARK_RGB

    current_part = None
    for pb in PLAYBOOKS:
        if pb["part"] != current_part:
            current_part = pb["part"]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(current_part)
            r.font.name = TITLE_FONT
            r.font.size = Pt(11)
            r.bold = True
            r.font.color.rgb = GOLD_RGB

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.2)
        r = p.add_run(f"#{pb['num']:02d}  \u2014  {pb['title']}")
        r.font.name = TITLE_FONT
        r.font.size = Pt(10.5)
        r.font.color.rgb = DARK_RGB

    # ==== PLAYBOOKS ====
    current_part = None
    for pb in PLAYBOOKS:
        # Part divider (every time the part changes)
        if pb["part"] != current_part:
            current_part = pb["part"]
            pp = doc.add_paragraph()
            set_page_break_before(pp)
            pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pp.paragraph_format.space_before = Pt(160)
            pp.paragraph_format.space_after = Pt(6)
            r = pp.add_run(f"PART {[p for p in 'IVX' if p][:1]}".rstrip())  # simple placeholder
            # Replace with proper roman
            pass

        # Skip the empty placeholder above and use real part title on page
        if pb["num"] in (1, 7, 11, 15, 18):
            pp = doc.add_paragraph()
            if pb["num"] != 1:
                set_page_break_before(pp)
            pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pp.paragraph_format.space_before = Pt(140 if pb["num"] != 1 else 40)
            pp.paragraph_format.space_after = Pt(6)
            r = pp.add_run(f"// PART //")
            r.font.name = TITLE_FONT
            r.font.size = Pt(11)
            r.font.color.rgb = GOLD_RGB

            pt = doc.add_paragraph()
            pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pt.paragraph_format.space_after = Pt(28)
            r = pt.add_run(pb["part"])
            r.font.name = TITLE_FONT
            r.font.size = Pt(24)
            r.bold = True
            r.font.color.rgb = DARK_RGB

        # Playbook header
        ph = doc.add_paragraph()
        if pb["num"] not in (1, 7, 11, 15, 18):
            set_page_break_before(ph)
        ph.paragraph_format.space_before = Pt(12 if pb["num"] in (1, 7, 11, 15, 18) else 0)
        ph.paragraph_format.space_after = Pt(4)
        r = ph.add_run(f"PLAYBOOK #{pb['num']:02d}")
        r.font.name = TITLE_FONT
        r.font.size = Pt(11)
        r.bold = True
        r.font.color.rgb = GOLD_RGB

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(pb["title"])
        r.font.name = TITLE_FONT
        r.font.size = Pt(18)
        r.bold = True
        r.font.color.rgb = DARK_RGB

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(14)
        r = p.add_run(pb["hook"])
        r.font.name = TITLE_FONT
        r.font.size = Pt(11)
        r.italic = True
        r.font.color.rgb = BROWN_RGB

        def _section(label, content, italic=False):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            r = p.add_run(label)
            r.font.name = TITLE_FONT
            r.font.size = Pt(9.5)
            r.bold = True
            r.font.color.rgb = GOLD_RGB

            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.3
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r = p.add_run(content)
            r.font.name = BODY_FONT
            r.font.size = Pt(10.5)
            if italic:
                r.italic = True
            r.font.color.rgb = DARK_RGB

        _section("THE STAKES", pb["stakes"])
        _section("THE DYNAMICS", pb["dynamics"])

        # 3 tactics
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        r = p.add_run("YOUR 3 TACTICS")
        r.font.name = TITLE_FONT
        r.font.size = Pt(9.5)
        r.bold = True
        r.font.color.rgb = GOLD_RGB

        for tname, tref in pb["tactics"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Inches(0.2)
            r1 = p.add_run("\u2022  ")
            r1.font.name = BODY_FONT
            r1.font.size = Pt(10.5)
            r1.bold = True
            r1.font.color.rgb = GOLD_RGB
            r2 = p.add_run(f"{tname}  ")
            r2.font.name = BODY_FONT
            r2.font.size = Pt(10.5)
            r2.bold = True
            r2.font.color.rgb = DARK_RGB
            r3 = p.add_run(f"\u2014  {tref}")
            r3.font.name = BODY_FONT
            r3.font.size = Pt(9.5)
            r3.italic = True
            r3.font.color.rgb = GRAY_RGB

        _section("OPENING LINE (WORD-FOR-WORD)", pb["opening"], italic=True)

        # Moves
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        r = p.add_run("MID-CONVERSATION MOVES")
        r.font.name = TITLE_FONT
        r.font.size = Pt(9.5)
        r.bold = True
        r.font.color.rgb = GOLD_RGB

        for mv in pb["moves"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.15)
            p.paragraph_format.line_spacing = 1.3
            r1 = p.add_run("\u2022  ")
            r1.font.name = BODY_FONT
            r1.font.size = Pt(10.5)
            r1.bold = True
            r1.font.color.rgb = GOLD_RGB
            r2 = p.add_run(mv)
            r2.font.name = BODY_FONT
            r2.font.size = Pt(10.5)
            r2.font.color.rgb = DARK_RGB

        # Objections
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        r = p.add_run("THE 2 OBJECTIONS YOU\u2019LL HIT")
        r.font.name = TITLE_FONT
        r.font.size = Pt(9.5)
        r.bold = True
        r.font.color.rgb = GOLD_RGB

        for obj, response in pb["objections"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Inches(0.2)
            r1 = p.add_run("They say:  ")
            r1.font.name = BODY_FONT
            r1.font.size = Pt(10)
            r1.bold = True
            r1.font.color.rgb = BROWN_RGB
            r2 = p.add_run(obj)
            r2.font.name = BODY_FONT
            r2.font.size = Pt(10)
            r2.italic = True
            r2.font.color.rgb = DARK_RGB

            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.left_indent = Inches(0.2)
            r1 = p.add_run("You respond:  ")
            r1.font.name = BODY_FONT
            r1.font.size = Pt(10)
            r1.bold = True
            r1.font.color.rgb = GOLD_RGB
            r2 = p.add_run(response)
            r2.font.name = BODY_FONT
            r2.font.size = Pt(10)
            r2.font.color.rgb = DARK_RGB

        _section("YOUR EXIT MOVE", pb["exit"], italic=True)
        _section("ONE MISTAKE TO AVOID", pb["mistake"])
        _section("THE DEBRIEF QUESTION (Fill Out Within 2 Hours)", pb["debrief"], italic=True)

    # ==== CLOSING ====
    p = doc.add_paragraph()
    set_page_break_before(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(180)
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("20 PLAYBOOKS  //  20 CONVERSATIONS  //  ONE LIFE")
    r.font.name = TITLE_FONT
    r.font.size = Pt(18)
    r.bold = True
    r.font.color.rgb = DARK_RGB

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("Pair each playbook with The Pre-Conversation Briefing.\nPull up the one you need. Fill it out in 10 minutes.\nWalk in prepared. Walk out with what you came for.")
    r.font.name = TITLE_FONT
    r.font.size = Pt(12)
    r.italic = True
    r.font.color.rgb = BROWN_RGB

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("\u2014  Nate Harlan")
    r.font.name = TITLE_FONT
    r.font.size = Pt(11)
    r.font.color.rgb = GRAY_RGB

    DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(DOCX_OUT))
    print(f"Docx saved: {DOCX_OUT}")
    print(f"Size: {DOCX_OUT.stat().st_size} bytes")
    print(f"Playbooks: {len(PLAYBOOKS)}")


if __name__ == "__main__":
    build_cover()
    build_docx()
