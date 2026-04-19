#!/usr/bin/env python3
"""
Bonus #4: The Reactance Detector Cheat Sheet

The capstone bonus. Ties directly to the book's core concept
(the Persuasion Detector and reactance theory).

Part 1: 24 phrase swaps across 4 categories. For each pair:
  - INSTEAD OF (the phrase that fires the listener's detector)
  - SAY (the quiet version that does the same work)
  - WHY (one-line mechanism note)

Part 2: 10 signals YOUR OWN detector is firing. Somebody is
running a tactic on you. What to notice and what to do.

Designed as a reference tool: readers scan it before a conversation
to eliminate their own detector-firing phrases, and pull it up
after a conversation when something felt off.
"""

from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REPO = Path(__file__).resolve().parents[2]
LOGO = REPO / "public" / "logo.png"
FONT_TTF = REPO / "funnel" / "assets" / "SpecialElite-Regular.ttf"
COVER_OUT = REPO / "funnel" / "assets" / "cover-bonus-4.png"
DOCX_OUT = REPO / "funnel" / "exports" / "bonus-4-reactance-detector-cheatsheet.docx"

BODY_FONT = "Georgia"
TITLE_FONT = "Georgia"

CREAM = (0xF4, 0xEC, 0xD8)
DARK = (0x1A, 0x1A, 0x1A)
BROWN = (0x5C, 0x3A, 0x1E)
GOLD = (0xD4, 0xA0, 0x17)

DARK_RGB = RGBColor(0x1A, 0x1A, 0x1A)
BROWN_RGB = RGBColor(0x5C, 0x3A, 0x1E)
GOLD_RGB = RGBColor(0xD4, 0xA0, 0x17)
RED_RGB = RGBColor(0x8B, 0x3D, 0x3D)
GREEN_RGB = RGBColor(0x3D, 0x8B, 0x5E)
GRAY_RGB = RGBColor(0x6B, 0x5B, 0x3E)


# ========== COVER ==========

def build_cover():
    WIDTH, HEIGHT = 1800, 2700
    img = Image.new("RGB", (WIDTH, HEIGHT), CREAM)
    draw = ImageDraw.Draw(img)

    label_font = ImageFont.truetype(str(FONT_TTF), 40)
    title_font = ImageFont.truetype(str(FONT_TTF), 110)
    subtitle_font = ImageFont.truetype(str(FONT_TTF), 58)
    author_font = ImageFont.truetype(str(FONT_TTF), 52)

    logo = Image.open(str(LOGO)).convert("RGBA")
    scale = 900 / logo.width
    logo_resized = logo.resize((int(logo.width * scale), int(logo.height * scale)), Image.LANCZOS)
    logo_x = (WIDTH - logo_resized.width) // 2
    logo_y = 380
    img.paste(logo_resized, (logo_x, logo_y), logo_resized)

    label_text = "// FREE BONUS  04 //"
    bbox = draw.textbbox((0, 0), label_text, font=label_font)
    label_y = logo_y + logo_resized.height + 80
    draw.text(((WIDTH - (bbox[2] - bbox[0])) // 2, label_y), label_text, fill=BROWN, font=label_font)

    title_y = label_y + 120
    title_line_spacing = 135

    title_lines = ["THE REACTANCE", "DETECTOR", "CHEATSHEET"]
    for i, line in enumerate(title_lines):
        bbox = draw.textbbox((0, 0), line, font=title_font)
        draw.text(((WIDTH - (bbox[2] - bbox[0])) // 2, title_y + i * title_line_spacing),
                  line, fill=DARK, font=title_font)

    rule_y = title_y + (len(title_lines) * title_line_spacing) + 60
    draw.rectangle([(WIDTH // 2 - 200, rule_y), (WIDTH // 2 + 200, rule_y + 3)], fill=GOLD)

    subtitle_lines = [
        "The Phrases That Fire",
        "The Detector,",
        "And What To Say Instead.",
    ]
    sub_y = rule_y + 80
    for line in subtitle_lines:
        bbox = draw.textbbox((0, 0), line, font=subtitle_font)
        draw.text(((WIDTH - (bbox[2] - bbox[0])) // 2, sub_y), line, fill=BROWN, font=subtitle_font)
        sub_y += 90

    author_text = "NATE HARLAN"
    bbox = draw.textbbox((0, 0), author_text, font=author_font)
    draw.text(((WIDTH - (bbox[2] - bbox[0])) // 2, HEIGHT - 280), author_text, fill=DARK, font=author_font)

    draw.rectangle([(60, 60), (WIDTH - 60, HEIGHT - 60)], outline=BROWN, width=2)

    COVER_OUT.parent.mkdir(parents=True, exist_ok=True)
    img.save(str(COVER_OUT), "PNG", dpi=(300, 300))
    print(f"Cover saved: {COVER_OUT}")


# ========== DOCX HELPERS ==========

def add_page_number(paragraph):
    run = paragraph.add_run()
    el = OxmlElement("w:fldChar")
    el.set(qn("w:fldCharType"), "begin")
    run._r.append(el)
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = "PAGE"
    run._r.append(it)
    el = OxmlElement("w:fldChar")
    el.set(qn("w:fldCharType"), "end")
    run._r.append(el)
    run.font.name = BODY_FONT
    run.font.size = Pt(10)
    run.font.color.rgb = BROWN_RGB


def set_page_break_before(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(OxmlElement("w:pageBreakBefore"))


# ========== CONTENT ==========

# Phrase swaps organized by category
# Tuple: (instead_of_phrase, say_instead, why_it_fires)
PHRASE_SWAPS = [
    {
        "category": "OPENING PHRASES",
        "subtitle": "The words that signal 'I am about to sell you something' in the first five seconds.",
        "swaps": [
            ("Thanks so much for taking the time.",
             "[Skip it. Open with a specific observation about them or the situation.]",
             "Every sales training starts with this exact phrase. Every sophisticated listener flags it instantly as the pre-pitch opener."),
            ("I was hoping we could...",
             "I wanted to...",
             "\"Hoping\" telegraphs weakness and pre-apologizes for making the ask. The listener registers that you don't think you have the right to be asking."),
            ("I know you're busy, but...",
             "[Skip the preamble. Get to the ask.]",
             "Everyone is busy. Stating it aloud signals you know you're low on their priority list. It invites them to say no before you've asked."),
            ("Can I ask you something?",
             "[Just ask.]",
             "Asking permission to ask is two asks. Both can be refused. The second ask was the one you cared about."),
            ("If you have a minute...",
             "[Just ask. If they don't have a minute, they'll say so.]",
             "Permission-seeking for basic attention reads as low-status and invites dismissal."),
            ("So, tell me about yourself.",
             "A specific question about something you already researched about them.",
             "This is the line every interviewer and networking event uses. It signals zero effort and zero genuine interest."),
        ],
    },
    {
        "category": "PITCH AND SALES LANGUAGE",
        "subtitle": "Phrases that announce you are about to do exactly what they're braced against.",
        "swaps": [
            ("To be honest with you...",
             "[Just say the honest thing.]",
             "Prefacing a statement with \"to be honest\" implies the previous things you said may not have been."),
            ("Just to be transparent...",
             "[Same. Just be.]",
             "Same problem. The phrase creates a distinction between your transparent moments and everything else."),
            ("Does that make sense?",
             "[Let the silence sit. If they have a question, they'll ask.]",
             "The universal sales tell for \"please validate that I am selling to you well.\" Every prospect has heard it a thousand times."),
            ("I'm going to shoot straight with you...",
             "[Shoot straight.]",
             "Announcing straightforwardness implies a previous indirectness. Real straightness is just a direct sentence."),
            ("Can I be real with you for a second?",
             "[Be real.]",
             "Same pattern. A real moment that needs labeling wasn't really real."),
            ("At the end of the day...",
             "[Skip it. State the point.]",
             "Filler phrase. Every listener's brain tunes out on the first syllable. It is a verbal runway that tells them to stop paying attention."),
        ],
    },
    {
        "category": "WEAK SOFTENERS",
        "subtitle": "Individual words and phrases that quietly downgrade whatever you say next.",
        "swaps": [
            ("I just wanted to...",
             "I wanted to...",
             "\"Just\" is the single most damaging word in professional English. It signals the thing that follows isn't important."),
            ("Actually, I think...",
             "I think...",
             "\"Actually\" signals disagreement or correction, triggering defensive posture before you've made your point."),
            ("Sorry to bother you...",
             "[Just ask.]",
             "If the question is worth asking, there is no need to apologize for it. If it isn't worth asking, don't ask."),
            ("I don't mean to be pushy but...",
             "[Just ask. Naming it won't help.]",
             "\"I don't mean to be X but\" is the setup for being X anyway. Listeners know the pattern. Naming it telegraphs what's coming."),
            ("I hope this isn't too much...",
             "[Make the request directly. Let them decide.]",
             "Pre-apologizing for the size of your ask makes them notice the size. It primes them to think yes, it is too much."),
            ("Would you mind if I...",
             "I'm going to [do the thing]. Just wanted to give you a heads-up.",
             "Asking permission is a request that can be refused. Stating an action with a courtesy notification is not."),
        ],
    },
    {
        "category": "QUESTIONS THAT ARE ACTUALLY PITCHES",
        "subtitle": "Soft-sell questions that every listener over twenty-five has been trained to spot.",
        "swaps": [
            ("Would you be open to...?",
             "What do you think about...?",
             "\"Would you be open to\" is the soft-sell structure every sales training program teaches. Pattern-matched instantly."),
            ("Does it make sense to...?",
             "How would that work on your side?",
             "\"Does it make sense\" telegraphs that you've already decided the answer. The real question is how would this actually work."),
            ("Don't you think we should...?",
             "What's your read on...?",
             "A leading question that pushes the listener toward agreement, which they resist by reflex. An open read-question invites them in."),
            ("Wouldn't it be great if...?",
             "[Skip it. State the actual idea.]",
             "Rhetorical \"wouldn't it be great\" language is a transparent pitch setup. The listener knows you're about to propose something."),
            ("Isn't it true that...?",
             "My read on this is...",
             "Demanding agreement with a premise is confrontational by design. Stating your read lets them respond without disputing a framing."),
            ("Can you see how...?",
             "Walk me through what you're seeing.",
             "\"Can you see how\" implies the listener has failed to see something obvious. It triggers defensiveness. Asking them to walk you through makes them the expert."),
        ],
    },
]


# Signals your own detector is firing (someone is running a tactic on you)
SELF_SIGNALS = [
    ("You feel a sudden urge to agree with whatever they just said.",
     "They have reached the end of a carefully constructed sequence of micro-yeses. The urge to agree is the payoff of commitment-and-consistency priming.",
     "Pause. Disagree out loud with something small, just to test whether you can."),

    ("You feel guilty for no clear reason.",
     "They did something for you (solicited or not) and the guilt is reciprocity pressure.",
     "Ask yourself whether the favor was something you actually wanted or needed. If not, the favor was a setup."),

    ("They are being unusually warm, flattering, or complimentary.",
     "Love bombing, or the start of a sales sequence. Normal warm people do not escalate warmth during a request.",
     "Slow the pace. Watch what happens when you say no to something small."),

    ("You notice they are repeating your last few words back to you.",
     "They are running Voss's mirror. The fact that you noticed it means it is being deployed visibly.",
     "Ask a clarifying question that requires their own words. \"What exactly do you mean by...?\""),

    ("They are naming what you are feeling, out loud.",
     "Labeling. Not inherently bad, but when you notice it happening, the label usually wasn't quite right.",
     "Correct the label out loud. The conversation moves into more honest territory."),

    ("You feel like you have to decide right now.",
     "Manufactured urgency. Most real deadlines are flexible by more than the pressure suggests.",
     "Ask yourself: is there an actual external deadline, or just their urgency? If the latter, say \"I'll let you know by [48 hours from now].\""),

    ("They're asking you a lot of questions, but you don't feel heard.",
     "Calibrated questions being used as a sales funnel, not actual curiosity.",
     "Stop answering. Ask one of your own: \"What are you trying to figure out?\""),

    ("You find yourself explaining something you weren't asked to explain.",
     "Silence pressure. They left a pause; your brain filled it. This is almost always a leak.",
     "Stop mid-sentence if you can. The over-explanation is usually the information they were trying to extract."),

    ("They're telling you what \"everyone\" thinks or does.",
     "Social proof manipulation. The \"everyone\" is often just them.",
     "Ask them to name one specific person. If they can't, you have your answer."),

    ("You leave the conversation feeling good but unsure what you agreed to.",
     "Charm was used to smooth over a commitment they wanted you to make.",
     "Go back to your notes, your email, your memory. Reconstruct the actual agreement. Then decide whether you would have agreed if you'd been less charmed."),
]


# ========== DOCX BUILD ==========

def build_docx():
    doc = Document()

    # ==== COVER ====
    cover_section = doc.sections[0]
    cover_section.page_width = Inches(6)
    cover_section.page_height = Inches(9)
    cover_section.top_margin = Inches(0)
    cover_section.bottom_margin = Inches(0)
    cover_section.left_margin = Inches(0)
    cover_section.right_margin = Inches(0)
    cover_section.header_distance = Inches(0)
    cover_section.footer_distance = Inches(0)

    cover_p = doc.add_paragraph()
    cover_p.paragraph_format.space_before = Pt(0)
    cover_p.paragraph_format.space_after = Pt(0)
    cover_p.paragraph_format.line_spacing = 1.0
    run = cover_p.add_run()
    run.add_picture(str(COVER_OUT), width=Inches(6), height=Inches(9))

    # ==== BODY ====
    body = doc.add_section(WD_SECTION_START.NEW_PAGE)
    body.page_width = Inches(6)
    body.page_height = Inches(9)
    body.top_margin = Inches(1.1)
    body.bottom_margin = Inches(0.9)
    body.left_margin = Inches(0.75)
    body.right_margin = Inches(0.75)
    body.header_distance = Inches(0.3)
    body.footer_distance = Inches(0.4)
    body.header.is_linked_to_previous = False
    body.footer.is_linked_to_previous = False

    h = body.header.paragraphs[0]
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hrun = h.add_run()
    hrun.add_picture(str(LOGO), width=Inches(0.9))

    f = body.footer.paragraphs[0]
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(f)

    # ==== INTRO ====
    intro_title = doc.add_paragraph()
    intro_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro_title.paragraph_format.space_before = Pt(12)
    intro_title.paragraph_format.space_after = Pt(24)
    r = intro_title.add_run("HOW TO USE THIS GUIDE")
    r.font.name = TITLE_FONT
    r.font.size = Pt(20)
    r.bold = True
    r.font.color.rgb = DARK_RGB

    intro_paragraphs = [
        "The Persuasion Detector is the automatic mental process that fires in a listener's head the moment they sense someone is trying to influence them. When it fires, they stop evaluating your message on its merits and start resisting it by reflex.",
        "Sixty years of psychology research says the same thing: the harder you push, the harder they push back. The skill this guide teaches is how to say the same thing in a different way, so the detector never fires.",
        "Part One is twenty-four phrase swaps across four categories. Each is a phrase that fires the listener's detector, and the quieter version that does the same work invisibly. Use it before your next important conversation: scan the list, eliminate the phrases you've been using, replace them with the alternatives.",
        "Part Two flips the direction. It is ten signals that somebody is running a tactic on you, right now. When you notice any of them, you have thirty seconds of useful information about what's actually happening in the conversation. The guide tells you what the signal means and what to do.",
        "This is a reference. Not a textbook. Come back to it the week before any conversation that matters.",
    ]
    for text in intro_paragraphs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.4
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.first_line_indent = Inches(0.2)
        r = p.add_run(text)
        r.font.name = BODY_FONT
        r.font.size = Pt(11)
        r.font.color.rgb = DARK_RGB

    # ==== PART 1 TITLE PAGE ====
    p = doc.add_paragraph()
    set_page_break_before(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("PART 1")
    r.font.name = TITLE_FONT
    r.font.size = Pt(14)
    r.bold = True
    r.font.color.rgb = BROWN_RGB

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("PHRASES THAT FIRE THE DETECTOR")
    r.font.name = TITLE_FONT
    r.font.size = Pt(18)
    r.bold = True
    r.font.color.rgb = DARK_RGB

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    r = p.add_run("Twenty-four phrase swaps across four categories. On the left: the phrase that fires the detector. On the right: the version that does the same work invisibly.")
    r.font.name = TITLE_FONT
    r.font.size = Pt(11)
    r.italic = True
    r.font.color.rgb = GRAY_RGB

    # ==== PHRASE SWAPS ====
    swap_number = 0
    for category in PHRASE_SWAPS:
        # Category header
        cat_p = doc.add_paragraph()
        set_page_break_before(cat_p)
        cat_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cat_p.paragraph_format.space_before = Pt(18)
        cat_p.paragraph_format.space_after = Pt(4)
        r = cat_p.add_run(category["category"])
        r.font.name = TITLE_FONT
        r.font.size = Pt(16)
        r.bold = True
        r.font.color.rgb = DARK_RGB

        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_p.paragraph_format.space_after = Pt(18)
        sub_p.paragraph_format.left_indent = Inches(0.3)
        sub_p.paragraph_format.right_indent = Inches(0.3)
        r = sub_p.add_run(category["subtitle"])
        r.font.name = TITLE_FONT
        r.font.size = Pt(10)
        r.italic = True
        r.font.color.rgb = GRAY_RGB

        for instead_of, say, why in category["swaps"]:
            swap_number += 1

            # Number + "INSTEAD OF"
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            r1 = p.add_run(f"{swap_number:02d}.  INSTEAD OF:   ")
            r1.font.name = TITLE_FONT
            r1.font.size = Pt(9.5)
            r1.bold = True
            r1.font.color.rgb = RED_RGB
            r2 = p.add_run(f"\u201C{instead_of}\u201D")
            r2.font.name = BODY_FONT
            r2.font.size = Pt(10.5)
            r2.italic = True
            r2.font.color.rgb = DARK_RGB

            # SAY
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.keep_with_next = True
            r1 = p.add_run("SAY:   ")
            r1.font.name = TITLE_FONT
            r1.font.size = Pt(9.5)
            r1.bold = True
            r1.font.color.rgb = GREEN_RGB
            r2 = p.add_run(say)
            r2.font.name = BODY_FONT
            r2.font.size = Pt(10.5)
            r2.italic = True
            r2.font.color.rgb = DARK_RGB

            # WHY
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.left_indent = Inches(0.3)
            r1 = p.add_run("WHY IT FIRES:   ")
            r1.font.name = TITLE_FONT
            r1.font.size = Pt(8.5)
            r1.bold = True
            r1.font.color.rgb = BROWN_RGB
            r2 = p.add_run(why)
            r2.font.name = BODY_FONT
            r2.font.size = Pt(9.5)
            r2.font.color.rgb = DARK_RGB

    # ==== PART 2 TITLE PAGE ====
    p = doc.add_paragraph()
    set_page_break_before(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("PART 2")
    r.font.name = TITLE_FONT
    r.font.size = Pt(14)
    r.bold = True
    r.font.color.rgb = BROWN_RGB

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("WHEN YOUR OWN DETECTOR FIRES")
    r.font.name = TITLE_FONT
    r.font.size = Pt(18)
    r.bold = True
    r.font.color.rgb = DARK_RGB

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run("Ten signals that somebody is running a tactic on you, right now. What each one means, and what to do about it.")
    r.font.name = TITLE_FONT
    r.font.size = Pt(11)
    r.italic = True
    r.font.color.rgb = GRAY_RGB

    # ==== SELF SIGNALS ====
    for i, (signal, meaning, do) in enumerate(SELF_SIGNALS, start=1):
        # Signal
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.keep_with_next = True
        r1 = p.add_run(f"{i:02d}.   ")
        r1.font.name = TITLE_FONT
        r1.font.size = Pt(11)
        r1.bold = True
        r1.font.color.rgb = GOLD_RGB
        r2 = p.add_run(signal)
        r2.font.name = TITLE_FONT
        r2.font.size = Pt(11.5)
        r2.bold = True
        r2.font.color.rgb = DARK_RGB

        # Meaning
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.keep_with_next = True
        r1 = p.add_run("What it means:   ")
        r1.font.name = TITLE_FONT
        r1.font.size = Pt(9)
        r1.bold = True
        r1.font.color.rgb = BROWN_RGB
        r2 = p.add_run(meaning)
        r2.font.name = BODY_FONT
        r2.font.size = Pt(10)
        r2.font.color.rgb = DARK_RGB

        # What to do
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Inches(0.35)
        r1 = p.add_run("What to do:   ")
        r1.font.name = TITLE_FONT
        r1.font.size = Pt(9)
        r1.bold = True
        r1.font.color.rgb = GREEN_RGB
        r2 = p.add_run(do)
        r2.font.name = BODY_FONT
        r2.font.size = Pt(10)
        r2.font.color.rgb = DARK_RGB

    # ==== CLOSING ====
    close_p = doc.add_paragraph()
    set_page_break_before(close_p)
    close_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    close_p.paragraph_format.space_before = Pt(60)
    close_p.paragraph_format.space_after = Pt(24)
    r = close_p.add_run("ONE LAST NOTE")
    r.font.name = TITLE_FONT
    r.font.size = Pt(20)
    r.bold = True
    r.font.color.rgb = DARK_RGB

    closing_paragraphs = [
        "The phrases in Part One are not a complete list. Your own speech contains others. Over the next month, pay attention to your own openers, your own softeners, and the places where you catch yourself using a phrase right out of a sales playbook. Add them to your personal version of this list.",
        "The signals in Part Two work the same way. Your own detector is better calibrated than you think. The ten in this guide are the most common. You will notice others in your own conversations once you start looking.",
        "The goal is not to learn a new vocabulary. The goal is to speak in a way that doesn't announce itself as persuasion, so the content of what you actually mean can be heard.",
        "For the full system behind this cheatsheet, see Shadow Persuasion: The 47 Counterintuitive Conversation Tactics That Make People Say Yes Without Realizing Why.",
    ]
    for text in closing_paragraphs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.4
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.first_line_indent = Inches(0.2)
        r = p.add_run(text)
        r.font.name = BODY_FONT
        r.font.size = Pt(11)
        r.font.color.rgb = DARK_RGB

    DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(DOCX_OUT))
    print(f"Docx saved: {DOCX_OUT}")
    print(f"Size: {DOCX_OUT.stat().st_size} bytes")
    print(f"Phrase swaps: {swap_number}")
    print(f"Self-detector signals: {len(SELF_SIGNALS)}")


if __name__ == "__main__":
    build_cover()
    build_docx()
