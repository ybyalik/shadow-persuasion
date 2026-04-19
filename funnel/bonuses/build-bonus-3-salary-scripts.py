#!/usr/bin/env python3
"""
Bonus #3: 48 Salary Negotiation Scripts

Word-for-word scripts for every phase of a salary conversation,
every common objection, and every special situation.

Organized into 6 parts:
  Part 1: Opening Moves (8 scripts)
  Part 2: Stating Your Number (6 scripts)
  Part 3: Handling Objections (15 scripts)
  Part 4: Negotiating the Offer (10 scripts)
  Part 5: Follow-Up and Emails (5 scripts)
  Part 6: Special Situations (4 scripts)

Each script has:
  - Title
  - When to use (short italic caption)
  - What to say (the actual word-for-word script)
  - Why it works (brief mechanism note)
"""

from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REPO = Path(__file__).resolve().parents[2]
LOGO = REPO / "public" / "logo.png"
FONT_TTF = REPO / "funnel" / "assets" / "SpecialElite-Regular.ttf"
COVER_OUT = REPO / "funnel" / "assets" / "cover-bonus-3.png"
DOCX_OUT = REPO / "funnel" / "exports" / "bonus-3-salary-negotiation-scripts.docx"

BODY_FONT = "Georgia"
TITLE_FONT = "Georgia"

CREAM = (0xF4, 0xEC, 0xD8)
DARK = (0x1A, 0x1A, 0x1A)
BROWN = (0x5C, 0x3A, 0x1E)
GOLD = (0xD4, 0xA0, 0x17)

DARK_RGB = RGBColor(0x1A, 0x1A, 0x1A)
BROWN_RGB = RGBColor(0x5C, 0x3A, 0x1E)
GOLD_RGB = RGBColor(0xD4, 0xA0, 0x17)
GRAY_RGB = RGBColor(0x6B, 0x5B, 0x3E)


# ========== COVER ==========

def build_cover():
    WIDTH, HEIGHT = 1800, 2700
    img = Image.new("RGB", (WIDTH, HEIGHT), CREAM)
    draw = ImageDraw.Draw(img)

    label_font = ImageFont.truetype(str(FONT_TTF), 40)
    title_font = ImageFont.truetype(str(FONT_TTF), 110)
    subtitle_font = ImageFont.truetype(str(FONT_TTF), 58)
    author_font = ImageFont.truetype(str(FONT_TTF), 52)

    # Logo
    logo = Image.open(str(LOGO)).convert("RGBA")
    scale = 900 / logo.width
    logo_resized = logo.resize((int(logo.width * scale), int(logo.height * scale)), Image.LANCZOS)
    logo_x = (WIDTH - logo_resized.width) // 2
    logo_y = 380
    img.paste(logo_resized, (logo_x, logo_y), logo_resized)

    # Bonus label
    label_text = "// FREE BONUS  03 //"
    bbox = draw.textbbox((0, 0), label_text, font=label_font)
    label_y = logo_y + logo_resized.height + 80
    draw.text(((WIDTH - (bbox[2] - bbox[0])) // 2, label_y), label_text, fill=BROWN, font=label_font)

    # Title: "48 SALARY" / "NEGOTIATION" / "SCRIPTS"
    title_y = label_y + 120
    title_line_spacing = 135

    title_lines = ["48 SALARY", "NEGOTIATION", "SCRIPTS"]
    for i, line in enumerate(title_lines):
        bbox = draw.textbbox((0, 0), line, font=title_font)
        draw.text(((WIDTH - (bbox[2] - bbox[0])) // 2, title_y + i * title_line_spacing),
                  line, fill=DARK, font=title_font)

    # Gold rule
    rule_y = title_y + (len(title_lines) * title_line_spacing) + 60
    draw.rectangle([(WIDTH // 2 - 200, rule_y), (WIDTH // 2 + 200, rule_y + 3)], fill=GOLD)

    # Subtitle
    subtitle_lines = [
        "Every Objection.",
        "Every Counter.",
        "Every Framing.",
    ]
    sub_y = rule_y + 80
    for line in subtitle_lines:
        bbox = draw.textbbox((0, 0), line, font=subtitle_font)
        draw.text(((WIDTH - (bbox[2] - bbox[0])) // 2, sub_y), line, fill=BROWN, font=subtitle_font)
        sub_y += 90

    # Author
    author_text = "NATE HARLAN"
    bbox = draw.textbbox((0, 0), author_text, font=author_font)
    draw.text(((WIDTH - (bbox[2] - bbox[0])) // 2, HEIGHT - 280), author_text, fill=DARK, font=author_font)

    # Border
    draw.rectangle([(60, 60), (WIDTH - 60, HEIGHT - 60)], outline=BROWN, width=2)

    COVER_OUT.parent.mkdir(parents=True, exist_ok=True)
    img.save(str(COVER_OUT), "PNG", dpi=(300, 300))
    print(f"Cover saved: {COVER_OUT}")


# ========== DOCX HELPERS ==========

def add_page_number(paragraph):
    run = paragraph.add_run()
    el = OxmlElement("w:fldChar")
    el.set(qn("w:fldCharType"), "begin")
    run._r.append(el)
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = "PAGE"
    run._r.append(it)
    el = OxmlElement("w:fldChar")
    el.set(qn("w:fldCharType"), "end")
    run._r.append(el)
    run.font.name = BODY_FONT
    run.font.size = Pt(10)
    run.font.color.rgb = BROWN_RGB


def set_page_break_before(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(OxmlElement("w:pageBreakBefore"))


# ========== SCRIPT CONTENT ==========

PARTS = [
    {
        "title": "PART 1",
        "subtitle": "OPENING MOVES",
        "blurb": "How to start the conversation. The opening determines the rest.",
        "scripts": [
            ("Requesting the raise conversation with your direct boss",
             "You want a dedicated meeting, not a squeeze into your regular 1-on-1.",
             "I'd like to book thirty minutes sometime this week to talk about my compensation. It's not something I want to squeeze into our regular 1-on-1 — it deserves its own meeting. Does Thursday at 2pm work, or later in the week?",
             "Requesting a dedicated meeting signals this is a real conversation, not a passing request. Pre-announcing the topic prevents you from ambushing them with it mid-1-on-1."),

            ("Opening with a skip-level manager",
             "Your direct manager has been stalling for months.",
             "I've been trying to move my compensation conversation forward with [manager] for the past four months without much progress. I wanted to loop you in directly because I respect your time and didn't want to escalate without letting you know I was going to.",
             "Names the pattern without attacking your manager. Frames your escalation as a professional courtesy, not a complaint."),

            ("Opening after you've been passed over",
             "You were passed over for a promotion or raise you expected.",
             "I want to understand Monday's decision so I can recalibrate what the next six months look like. I'm not asking you to reverse it. I'm asking what I would need to do differently so we're not having this same conversation next cycle.",
             "Doesn't plead for reversal, which rarely works. Forces concrete criteria, which gives you a measurable goal to hit."),

            ("Opening during an existing performance review",
             "You're in a review where comp might not naturally come up.",
             "Before we wrap up the feedback discussion, I want to make sure we cover compensation. I know that's technically a separate conversation, but I'd rather link them while the context is fresh for both of us.",
             "Stakes a claim on the compensation topic before the meeting ends, preventing the 'we'll talk about that separately' deflection that goes nowhere."),

            ("Opening cold with no upcoming review",
             "There's no cycle. You just need to have the conversation.",
             "I don't want to wait until review season to have this conversation. I've been carrying it around for a few months, and I think it deserves a dedicated thirty minutes with you. Can we book it for next week?",
             "Naming that you've been carrying it signals seriousness. Asking for a dedicated meeting rather than an agenda squeeze shows you expect this to be substantive."),

            ("Opening via email to request the meeting",
             "You want the request in writing so there's a paper trail.",
             "[Manager], I'd like to schedule thirty minutes with you sometime in the next two weeks to discuss my compensation. I've put together some thoughts and want to share them with you directly. Would Tuesday the 12th at 10am work? If not, I can flex anytime between the 10th and the 22nd. Thanks, [Your name]",
             "Short, specific, scheduled. Gives options without seeming flexible to the point of weakness. 'Thoughts' signals preparation without previewing your case."),

            ("Opening when you have a competing offer",
             "You have a real offer and want leverage without sounding like a threat.",
             "I want to be straight with you. I received an offer from [Company] last week at [number]. I'm not shopping myself around and I'm not going to pretend I'm ready to leave over this. But I'd like to understand what's possible here before I make any decisions.",
             "Volunteering the information (before they discover it) builds trust. Framing yourself as not eager to leave removes the adversarial tone."),

            ("Opening a promotion conversation, not just a raise",
             "You want the title bump, which usually comes with the raise.",
             "I want to have a specific conversation today about promotion, not just compensation. My sense is that I've been doing the work of a [target level] for about nine months, and I'd like to make that official. I want to understand what's blocking it if it's not the next step.",
             "'Doing the work of' is the canonical promotion case. You're not asking for something new; you're asking them to acknowledge what's already true."),
        ],
    },
    {
        "title": "PART 2",
        "subtitle": "STATING YOUR NUMBER",
        "blurb": "How to put a number on the table without losing it in the first five minutes.",
        "scripts": [
            ("Stating a specific number first",
             "You know what you want and want to anchor.",
             "I've been thinking about this carefully, and the number I think is right is [specific number]. I can walk you through how I got there if you want. I wanted to put it on the table first so we can have a real conversation about it.",
             "Anchoring high with a specific number sets the range. Offering to explain afterward (instead of pre-justifying) projects confidence."),

            ("Responding when they ask \"what are you looking for?\"",
             "They've turned the question on you first.",
             "That's a fair question. Based on the work I've been doing and the market data for this role, [specific number] feels right to me. I'm open to a conversation about it.",
             "Doesn't dodge the question (dodging reads as weak). Gives a specific number anchored in two reasons. 'Open to a conversation' signals flexibility without rolling over."),

            ("Anchoring high without committing to a specific number",
             "You want to anchor the range but not hand them a number to reject.",
             "I've been looking at comparable roles at peer companies and the range I'm seeing is [high end of range] to [even higher]. I think I'm comfortably in that range. I want to understand what's realistic from your side before I nail down a specific ask.",
             "Anchors on the top of the market without committing to a number. Invites them to share their range, which gives you data."),

            ("Framing the number as market data",
             "You want to take the personal element out of the ask.",
             "I've pulled data from Levels.fyi, Glassdoor, and a few conversations with people in similar roles. Median for this role at a company our size is [number]. I'm not asking for above market. I'm asking for at market.",
             "Framing as 'at market' (not 'above') makes your request seem reasonable rather than aggressive. Specific sources make it harder to dismiss."),

            ("Framing the number as your replacement cost",
             "You want to quietly remind them what losing you would cost.",
             "If I were to leave and you had to replace me, you'd be looking at about six months of recruiting, three months of onboarding, and somewhere between [X] and [Y] in all-in cost to get someone to my level. What I'm asking for is significantly less than that.",
             "Moves the frame from 'how much you pay me' to 'how much it would cost to replace me.' The replacement number is almost always higher than the raise ask."),

            ("When you don't have a specific number yet",
             "You want to open the conversation without being pinned down.",
             "I want to be honest — I don't have a specific number yet. I'm at the point where I know my current comp doesn't reflect what I'm doing. I'd like to figure out together what the right number is.",
             "Transparency about not having a number, paired with the statement that current comp is wrong, reframes the conversation as collaborative rather than adversarial."),
        ],
    },
    {
        "title": "PART 3",
        "subtitle": "HANDLING OBJECTIONS",
        "blurb": "The fifteen objections you will actually hear, with the exact response to each.",
        "scripts": [
            ("\"It's not in the budget.\"",
             "The classic first-line deflection.",
             "I hear you. Budgets are built months in advance and they're usually tight. I want to ask two questions. First, when does the next budget cycle open? Second, if we can't move in this cycle, what would we need to agree on now so that when the next cycle opens, this is already handled?",
             "Doesn't fight the budget claim (which is usually true). Forces a concrete commitment for the next window."),

            ("\"Now isn't the right time.\"",
             "Vague timing deflection.",
             "Help me understand what the right time would look like. What has to be true for this conversation to move forward?",
             "A single calibrated question that refuses the vague deflection without pushing back. Forces them to describe concrete conditions, which you can then meet."),

            ("\"We need to wait until review cycle.\"",
             "They want to defer to a scheduled process.",
             "I understand the review cycle is the official process. I want to flag two concerns. First, that's six months away, and I don't want to feel underpaid for six more months. Second, review-cycle raises tend to be capped at a percentage. I'd rather have this conversation now, where we can land on the right number, than wait and get something formulaic.",
             "Acknowledges their position without capitulating. Names the real downside of the review process (formulaic, capped) without attacking the process itself."),

            ("\"I don't have the authority to approve that.\"",
             "They're deflecting to a higher-up.",
             "That's fair. Who does have the authority? I'd like to understand what the chain of approval looks like and whether there's anything you need from me to make the case to them.",
             "Doesn't try to push them into saying yes when they can't. Maps the decision-making path and positions your manager as your advocate rather than your obstacle."),

            ("\"Other people on the team are at the same level.\"",
             "Equity-based pushback.",
             "That's an important concern. Two things. First, I'd be curious how my performance is rated relative to theirs — I don't want special treatment, I want to be compensated for what I've contributed. Second, if we decide I'm performing at the same level as others, the real conversation is whether the whole team is underpaid, not whether I deserve more.",
             "Refuses the false equity frame. Either you're outperforming (deserve more) or you're not (then it's a team-wide issue, which is their problem to solve)."),

            ("\"Your performance hasn't justified that.\"",
             "They're critiquing your work as the reason to decline.",
             "I want to understand what would have justified it. What specifically did I not do this year that would have made this a different conversation? I'd like to leave this meeting with a concrete list.",
             "Doesn't defend (defending is losing). Extracts specific actionable feedback. If they can't give specifics, the critique was pretextual, and you now know it."),

            ("\"Let me see what I can do.\"",
             "The classic non-committal response.",
             "I appreciate that. Can we agree on a specific date when you'll come back to me with a concrete answer? I don't want to be the one chasing this for the next three months.",
             "'Let me see what I can do' is how raises die quietly. Pinning a specific date creates accountability. If they refuse to commit to a date, that is also information."),

            ("\"We have to be fair across the team.\"",
             "Fairness-based pushback.",
             "I agree with that principle. Fairness means people are paid for what they contribute. Can we compare what I've contributed to what the team's median performer has contributed, and see whether my current comp reflects that? If it does, I'll drop it. If it doesn't, we need to talk about it.",
             "Flips 'fairness' from a shield they're using against you into a framework that supports your case, assuming you're actually performing above median."),

            ("\"That's above market for this role.\"",
             "They're claiming your ask is too high.",
             "Can you share the data you're using? I've pulled from [three sources] and I'm seeing a different range. I'm happy to walk through my sources and reconcile with yours.",
             "Forces them to produce actual data, which they usually don't have. Most 'above market' claims are defensive rather than factual."),

            ("\"You just got a raise recently.\"",
             "They're pointing to a recent increase.",
             "I appreciate the raise I got [six months ago]. My point is that the scope of what I'm doing now has changed significantly since then — specifically [concrete examples]. A raise that made sense six months ago doesn't reflect what I'm doing today.",
             "Doesn't fight the prior raise. Anchors the discussion on scope change, which is the actual reason the number needs to move again."),

            ("\"We're in a hiring freeze.\"",
             "They're using broader company financial conditions.",
             "A hiring freeze doesn't affect retention of existing people. If anything, it makes retention more important — you can't replace me if I leave. I want to understand whether the freeze applies to compensation for current employees or whether it's specifically about headcount.",
             "Distinguishes between two things (hiring vs. retention) that they've conflated. The distinction usually works in your favor."),

            ("\"Show me a competing offer.\"",
             "They want you to prove your market value with another job offer.",
             "I'm not in a place where I want to shop myself around just to prove a point. I want to stay here, which is why I'm having this conversation with you. What I can do is share market data. What else would help you make the case?",
             "Refuses the 'bring me an offer' game, which always ends badly. Keeps the conversation about your retention rather than your leverage."),

            ("\"You're not ready yet.\"",
             "They're claiming you're not developmentally ready.",
             "Help me understand specifically what 'ready' looks like. I don't want a vague answer. I want a concrete list of what I need to demonstrate. If I do those things, are we agreed that the conversation we're having now becomes a yes?",
             "Forces concrete criteria. Most 'not ready' responses are actually 'not budgeted.' Closes the loop by getting agreement on what success looks like."),

            ("\"Can you wait until next quarter?\"",
             "They want to defer.",
             "I can wait a quarter if I know what we're agreeing to today. If we can agree now on the number that will be implemented next quarter and put it in writing, I'm happy to wait. If 'next quarter' means we revisit this conversation from scratch in three months, I'd rather push harder now.",
             "Distinguishes commitment-to-execute-later from soft defer-and-forget. Either they give you the commitment or they reveal the deferral for what it is."),

            ("\"That's more than I make.\"",
             "A rare but real objection from an underpaid manager.",
             "I hear you. I want to be respectful about that. What I'm asking for is based on market data and the scope of what I'm doing, not on what I think anyone else should or shouldn't make. I don't want to make this awkward — but I also can't adjust my ask based on what you're paid.",
             "Acknowledges the awkwardness without retreating. Separates your comp from your manager's, which is actually a separate issue."),
        ],
    },
    {
        "title": "PART 4",
        "subtitle": "NEGOTIATING THE OFFER",
        "blurb": "What to say after they come back with a number. Most of the real money moves in this phase.",
        "scripts": [
            ("When they offer less than you asked",
             "They've moved, but not enough.",
             "I appreciate the movement. That number doesn't quite get us to where I think it should be. What would need to be true on your end for you to get me to [your original number or close to it]?",
             "Acknowledges their effort without rejecting outright. Re-anchors on your number. 'What would need to be true' puts them in the work of figuring out how to get there."),

            ("When they offer exactly what you asked",
             "They said yes to the number. Don't close yet.",
             "Thank you. Let me think through whether there are other pieces I want to discuss while we're in this conversation — title, equity, timing, anything else we haven't touched. Would you be open to continuing once I've thought through it?",
             "The biggest mistake when they say yes is to immediately close. Holding the door open for other items often nets more than the salary alone."),

            ("When they offer above what you asked",
             "Rare but real. Don't leave money on the table.",
             "That's generous. Thank you. Let me make sure I understand the full offer — is the number what's in the job level, or did you build in some margin on top? I want to understand the ceiling so I know what to think about for next year.",
             "Don't act surprised, but get information about where the next conversation starts. The ceiling matters more than the current offer."),

            ("Negotiating total compensation, not just salary",
             "You want to expand the conversation beyond base.",
             "If base is tight, I'd like to talk about total compensation. Signing bonus, equity refresher, performance bonus target, vacation, anything else in your authority. Those are different budgets usually. What's flexible on your side?",
             "Often base is in a strict band while other comp elements have more flexibility. Explicitly listing options prompts consideration."),

            ("Negotiating equity or stock",
             "You want to push on the equity side specifically.",
             "Equity matters to me. If base is capped, I'd trade some of the base increase for an equity refresher at a meaningful multiple. What does the equity refresh picture look like for someone at my level?",
             "Many companies have underutilized equity refresh programs. Offering a trade (base for equity) signals you're a collaborative negotiator."),

            ("Negotiating a signing or retention bonus",
             "A one-time payment when base is capped.",
             "If annual base is where it needs to be but you want to keep cash conservative, a signing or retention bonus structured to vest over 12 or 18 months could work for both of us. I get the upfront, you get the retention leverage.",
             "Frames it as a win-win. These bonuses often come from different budgets than salary and are easier to approve."),

            ("Negotiating title",
             "You want the title bump regardless of salary.",
             "Separate from the salary conversation — I want to ask about title. I've been doing [current-level-plus] work for [timeframe] and I'd like the title to reflect that. Is that something we can land on in this conversation, or is that a separate process?",
             "Decouples title from salary. Titles often have more administrative flexibility than salaries."),

            ("Negotiating remote work",
             "You want location flexibility as part of the package.",
             "Non-monetary thing I want to put on the table: I'd like to formalize [X days a week remote / full remote / a specific arrangement]. If I can't get the full comp ask, this is something I'd value almost as much. What's the policy flexibility on your side?",
             "Explicitly naming non-monetary asks as part of the same conversation prevents them from being treated as separate, smaller conversations."),

            ("Negotiating start date",
             "You're accepting an offer and want more time before starting.",
             "I'm going to need four weeks before I can start, not two. I want to wrap up cleanly at my current role and take a week between. I want to make sure that's okay before we lock in the start date.",
             "States the ask directly, provides the reason (clean transition plus rest), and asks for confirmation rather than permission."),

            ("Getting it in writing",
             "They verbally agreed. Now you need documentation.",
             "Thanks for this. Can you send me an email today or tomorrow that captures what we agreed to? I'd rather have it in writing in case either of us forgets details.",
             "Frames written confirmation as mutual protection rather than distrust. 'In case either of us forgets' is diplomatic; the real reason is to prevent later renegotiation."),
        ],
    },
    {
        "title": "PART 5",
        "subtitle": "FOLLOW-UP AND EMAILS",
        "blurb": "The written word after the conversation. Shorter is always better.",
        "scripts": [
            ("Follow-up email on the same day",
             "Document the conversation while it's fresh.",
             "[Manager], thanks for the time today. To summarize what we discussed: you're going to [action] by [date], and I'm going to [action]. Let me know if I missed anything. Looking forward to the next step. [Your name]",
             "Short, specific, documents the commitments. Creates a paper trail without feeling like a gotcha."),

            ("Follow-up email after 3 days with no response",
             "They haven't gotten back to you.",
             "[Manager], following up on last Tuesday's conversation. I don't want to rush you — I know these things take time internally — but wanted to check in on where things stand. Happy to chat any time this week. [Your name]",
             "Low pressure, acknowledges internal process, doesn't create conflict. The check-in is its own signal that you're not dropping this."),

            ("Follow-up email after 1 week with no response",
             "A full week has passed.",
             "[Manager], it's been a week since our conversation and I haven't heard back. I want to make sure this hasn't fallen off the plate. What does the timeline look like from here?",
             "Slightly more direct than the 3-day version. 'Fallen off the plate' gives them a procedural excuse without accusing them of ignoring you."),

            ("Thank-you email after they agree",
             "They said yes to the ask.",
             "[Manager], thanks for the conversation and for making this happen. Looking forward to [specific next thing you're working on]. Let me know when the paperwork comes through. [Your name]",
             "Short, warm, immediately moves to what's next. Does not re-celebrate the raise, which would only remind them of the money; just transitions to the work."),

            ("Email after they say no, keeping the door open",
             "They declined. You're staying for now.",
             "[Manager], thanks for being straight with me about where things stand. I want to keep the conversation open for the next review cycle. In the meantime, I'd like to get clear on what you'd want to see from me between now and then, so we both know the bar. Can we put 30 minutes on the calendar for [next week] to go through that? [Your name]",
             "No sulking, no threats, immediately moves to 'what would need to be different next time.' Keeps you positioned as rational and methodical, not emotional."),
        ],
    },
    {
        "title": "PART 6",
        "subtitle": "SPECIAL SITUATIONS",
        "blurb": "The conversations that don't fit the normal pattern. Handle with care.",
        "scripts": [
            ("Raise request after a promotion without a raise",
             "You got the title but not the money.",
             "I appreciate the promotion. I want to name something directly — a promotion without a corresponding comp adjustment is just new responsibility. I want to understand when the comp piece catches up, because if the answer is 'it doesn't,' I need to think about that.",
             "Names the common trick (title bump without pay bump) directly. The final clause is a subtle but clear signal that you'll consider your options if they don't fix it."),

            ("When you're significantly underpaid versus market",
             "You've done the research. You know.",
             "I want to surface something that's been bothering me. Based on the data I've pulled, I'm paid about [X%] below market for my role. I'm not asking you to match the top of the market. I'm asking you to get me to at least the median. What's the path?",
             "Specific percentage is disarming. Asking for median (not top) is modest enough to be hard to refuse. 'What's the path' assumes they'll help rather than block."),

            ("Counter-offer when you've already accepted another job",
             "They're trying to retain you after you've signaled you're leaving.",
             "I want to be honest — if this comp had been on the table three months ago, I wouldn't be having the other conversation. I appreciate the offer, and I need 48 hours to think about it. I'll come back to you Friday.",
             "Names the timing issue (comp should have been proactive). Buys time to compare. Specific return date shows seriousness."),

            ("When your boss asks \"what would make you stay?\"",
             "They've asked the question outright.",
             "That's the right question. Before I answer, I want to be honest — the fact that you're asking after I've already signaled I'm thinking about leaving is part of the issue. I'd rather not have every comp conversation require me to almost quit to move it forward. So the comp piece is [specific number], and beyond that, I want us to agree on what a proactive review process looks like so we're not doing this again.",
             "Names the dysfunctional pattern (having to threaten to leave to get paid). Sets both a dollar number and a process change. Most retention offers fix the number but not the pattern."),
        ],
    },
]


# ========== DOCX BUILD ==========

def build_docx():
    doc = Document()

    # ==== COVER ====
    cover_section = doc.sections[0]
    cover_section.page_width = Inches(6)
    cover_section.page_height = Inches(9)
    cover_section.top_margin = Inches(0)
    cover_section.bottom_margin = Inches(0)
    cover_section.left_margin = Inches(0)
    cover_section.right_margin = Inches(0)
    cover_section.header_distance = Inches(0)
    cover_section.footer_distance = Inches(0)

    cover_p = doc.add_paragraph()
    cover_p.paragraph_format.space_before = Pt(0)
    cover_p.paragraph_format.space_after = Pt(0)
    cover_p.paragraph_format.line_spacing = 1.0
    run = cover_p.add_run()
    run.add_picture(str(COVER_OUT), width=Inches(6), height=Inches(9))

    # ==== BODY ====
    body = doc.add_section(WD_SECTION_START.NEW_PAGE)
    body.page_width = Inches(6)
    body.page_height = Inches(9)
    body.top_margin = Inches(1.1)
    body.bottom_margin = Inches(0.9)
    body.left_margin = Inches(0.75)
    body.right_margin = Inches(0.75)
    body.header_distance = Inches(0.3)
    body.footer_distance = Inches(0.4)
    body.header.is_linked_to_previous = False
    body.footer.is_linked_to_previous = False

    # Header
    h = body.header.paragraphs[0]
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hrun = h.add_run()
    hrun.add_picture(str(LOGO), width=Inches(0.9))

    # Footer
    f = body.footer.paragraphs[0]
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(f)

    # ==== INTRO ====
    intro_title = doc.add_paragraph()
    intro_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro_title.paragraph_format.space_before = Pt(12)
    intro_title.paragraph_format.space_after = Pt(24)
    r = intro_title.add_run("HOW TO USE THIS GUIDE")
    r.font.name = TITLE_FONT
    r.font.size = Pt(20)
    r.bold = True
    r.font.color.rgb = DARK_RGB

    intro_paragraphs = [
        "This guide contains forty-eight word-for-word scripts for the salary and compensation conversations you are going to have over the course of your career.",
        "The scripts are organized into six parts: opening moves, stating your number, handling objections, negotiating the offer once it lands, following up, and special situations. You will not read this cover to cover. You will pull up the specific script you need, twenty minutes before the conversation.",
        "Two things about the scripts themselves.",
        "First, the exact wording matters. Replacing a word with what you think is a better synonym usually neutralizes the move. If you find yourself wanting to soften a line, that softening is probably exactly the thing that makes the original line work.",
        "Second, every script has a short 'why it works' note underneath. The why matters more than the what. If you understand why a line lands, you can improvise when the actual conversation goes somewhere unexpected.",
        "The most common failure with scripts is memorizing them verbatim without absorbing the logic. Read the script. Read the why. Practice it out loud once. Then trust yourself in the conversation.",
    ]
    for text in intro_paragraphs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.4
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.first_line_indent = Inches(0.2)
        r = p.add_run(text)
        r.font.name = BODY_FONT
        r.font.size = Pt(11)
        r.font.color.rgb = DARK_RGB

    # ==== PARTS AND SCRIPTS ====
    script_number = 0
    for part in PARTS:
        # Part title page
        part_p = doc.add_paragraph()
        set_page_break_before(part_p)
        part_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        part_p.paragraph_format.space_before = Pt(60)
        part_p.paragraph_format.space_after = Pt(6)
        r = part_p.add_run(part["title"])
        r.font.name = TITLE_FONT
        r.font.size = Pt(14)
        r.bold = True
        r.font.color.rgb = BROWN_RGB

        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_p.paragraph_format.space_after = Pt(18)
        r = sub_p.add_run(part["subtitle"])
        r.font.name = TITLE_FONT
        r.font.size = Pt(22)
        r.bold = True
        r.font.color.rgb = DARK_RGB

        blurb_p = doc.add_paragraph()
        blurb_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        blurb_p.paragraph_format.space_after = Pt(36)
        blurb_p.paragraph_format.left_indent = Inches(0.4)
        blurb_p.paragraph_format.right_indent = Inches(0.4)
        r = blurb_p.add_run(part["blurb"])
        r.font.name = TITLE_FONT
        r.font.size = Pt(11)
        r.italic = True
        r.font.color.rgb = GRAY_RGB

        # Individual scripts
        for script in part["scripts"]:
            script_number += 1
            title, when_to_use, script_text, why = script

            # Script header
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            r = p.add_run(f"SCRIPT #{script_number:02d}.  {title}")
            r.font.name = TITLE_FONT
            r.font.size = Pt(11.5)
            r.bold = True
            r.font.color.rgb = DARK_RGB

            # When to use (italic gray)
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.left_indent = Inches(0.0)
            p.paragraph_format.keep_with_next = True
            r = p.add_run(f"When to use: {when_to_use}")
            r.font.name = BODY_FONT
            r.font.size = Pt(9.5)
            r.italic = True
            r.font.color.rgb = GRAY_RGB

            # Script text (blockquote style)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.right_indent = Inches(0.3)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.3
            p.paragraph_format.keep_with_next = True
            r1 = p.add_run("\u201C")
            r1.font.name = BODY_FONT
            r1.font.size = Pt(11)
            r1.bold = True
            r1.font.color.rgb = GOLD_RGB
            r2 = p.add_run(script_text)
            r2.font.name = BODY_FONT
            r2.font.size = Pt(11)
            r2.font.color.rgb = DARK_RGB
            r3 = p.add_run("\u201D")
            r3.font.name = BODY_FONT
            r3.font.size = Pt(11)
            r3.bold = True
            r3.font.color.rgb = GOLD_RGB

            # Why it works
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.0)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.25
            r1 = p.add_run("Why it works:  ")
            r1.font.name = BODY_FONT
            r1.font.size = Pt(9.5)
            r1.bold = True
            r1.font.color.rgb = BROWN_RGB
            r2 = p.add_run(why)
            r2.font.name = BODY_FONT
            r2.font.size = Pt(9.5)
            r2.font.color.rgb = DARK_RGB

    # ==== CLOSING ====
    close_p = doc.add_paragraph()
    set_page_break_before(close_p)
    close_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    close_p.paragraph_format.space_before = Pt(60)
    close_p.paragraph_format.space_after = Pt(24)
    r = close_p.add_run("ONE LAST NOTE")
    r.font.name = TITLE_FONT
    r.font.size = Pt(20)
    r.bold = True
    r.font.color.rgb = DARK_RGB

    closing_paragraphs = [
        "A salary negotiation is not a performance. It is a conversation in which you are the only person representing your interests. The scripts in this guide are designed to keep you in that position throughout.",
        "The biggest mistake people make in these conversations is not failing to execute a script. It is abandoning the conversation prematurely. Staying in the room, past the first no, past the second deflection, past the vague 'let me see what I can do,' is often what determines whether the number moves.",
        "You will forget these scripts in the heat of the moment. That is fine. If you only remember three things, remember these: ask a clarifying question instead of defending, pin every agreement to a specific date, and follow up in writing within twenty-four hours.",
        "For the full system behind these scripts, see Shadow Persuasion: The 47 Counterintuitive Conversation Tactics That Make People Say Yes Without Realizing Why.",
    ]
    for text in closing_paragraphs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.4
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.first_line_indent = Inches(0.2)
        r = p.add_run(text)
        r.font.name = BODY_FONT
        r.font.size = Pt(11)
        r.font.color.rgb = DARK_RGB

    DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(DOCX_OUT))
    print(f"Docx saved: {DOCX_OUT}")
    print(f"Size: {DOCX_OUT.stat().st_size} bytes")
    print(f"Scripts: {script_number}")


if __name__ == "__main__":
    build_cover()
    build_docx()
