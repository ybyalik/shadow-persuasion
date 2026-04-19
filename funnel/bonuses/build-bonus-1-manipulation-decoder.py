#!/usr/bin/env python3
"""
Bonus #1: The Manipulation Tactics Decoder

50 common manipulation tactics organized in 5 categories of 10:
- Emotional Manipulation
- Language Tricks
- Social Pressure
- Information Manipulation
- Power Plays

Each tactic has: Name, What You'll Hear, Red Flag, Counter.

Generates:
- cover-bonus-1.png (cream cover image matching book brand)
- bonus-1-manipulation-tactics-decoder.docx (full document)
"""

from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REPO = Path(__file__).resolve().parents[2]
LOGO = REPO / "public" / "logo.png"
FONT_TTF = REPO / "funnel" / "assets" / "SpecialElite-Regular.ttf"
COVER_OUT = REPO / "funnel" / "assets" / "cover-bonus-1.png"
DOCX_OUT = REPO / "funnel" / "exports" / "bonus-1-manipulation-tactics-decoder.docx"

BODY_FONT = "Georgia"
TITLE_FONT = "Georgia"

# Brand colors
CREAM = (0xF4, 0xEC, 0xD8)
DARK = (0x1A, 0x1A, 0x1A)
BROWN = (0x5C, 0x3A, 0x1E)
GOLD = (0xD4, 0xA0, 0x17)
RED = (0x8B, 0x3D, 0x3D)

DARK_RGB = RGBColor(0x1A, 0x1A, 0x1A)
BROWN_RGB = RGBColor(0x5C, 0x3A, 0x1E)
GOLD_RGB = RGBColor(0xD4, 0xA0, 0x17)
RED_RGB = RGBColor(0x8B, 0x3D, 0x3D)
GRAY_RGB = RGBColor(0x6B, 0x5B, 0x3E)

# ========== COVER IMAGE ==========

def build_cover():
    WIDTH, HEIGHT = 1800, 2700
    img = Image.new("RGB", (WIDTH, HEIGHT), CREAM)
    draw = ImageDraw.Draw(img)

    label_font = ImageFont.truetype(str(FONT_TTF), 40)
    title_font = ImageFont.truetype(str(FONT_TTF), 92)
    subtitle_font = ImageFont.truetype(str(FONT_TTF), 58)
    author_font = ImageFont.truetype(str(FONT_TTF), 44)

    # Logo at top
    logo = Image.open(str(LOGO)).convert("RGBA")
    target_w = 900
    scale = target_w / logo.width
    logo_resized = logo.resize((int(logo.width * scale), int(logo.height * scale)), Image.LANCZOS)
    img.paste(logo_resized, ((WIDTH - logo_resized.width) // 2, 360), logo_resized)

    # Bonus label
    label_text = "// FREE BONUS  01 //"
    bbox = draw.textbbox((0, 0), label_text, font=label_font)
    draw.text(((WIDTH - (bbox[2] - bbox[0])) // 2, 880), label_text, fill=BROWN, font=label_font)

    # Title
    title_text = "THE MANIPULATION"
    bbox = draw.textbbox((0, 0), title_text, font=title_font)
    draw.text(((WIDTH - (bbox[2] - bbox[0])) // 2, 1020), title_text, fill=DARK, font=title_font)

    title_text2 = "TACTICS DECODER"
    bbox = draw.textbbox((0, 0), title_text2, font=title_font)
    draw.text(((WIDTH - (bbox[2] - bbox[0])) // 2, 1140), title_text2, fill=DARK, font=title_font)

    # Gold rule
    draw.rectangle([(WIDTH // 2 - 180, 1310), (WIDTH // 2 + 180, 1313)], fill=GOLD)

    # Subtitle
    subtitle_lines = [
        "50 Common Manipulation Tactics",
        "How To Spot Them",
        "How To Shut Them Down",
    ]
    sub_y = 1380
    for line in subtitle_lines:
        bbox = draw.textbbox((0, 0), line, font=subtitle_font)
        draw.text(((WIDTH - (bbox[2] - bbox[0])) // 2, sub_y), line, fill=BROWN, font=subtitle_font)
        sub_y += 90

    # Author
    author_text = "NATE HARLAN"
    bbox = draw.textbbox((0, 0), author_text, font=author_font)
    draw.text(((WIDTH - (bbox[2] - bbox[0])) // 2, HEIGHT - 280), author_text, fill=DARK, font=author_font)

    # Border
    draw.rectangle([(60, 60), (WIDTH - 60, HEIGHT - 60)], outline=BROWN, width=2)

    COVER_OUT.parent.mkdir(parents=True, exist_ok=True)
    img.save(str(COVER_OUT), "PNG", dpi=(300, 300))
    print(f"Cover saved: {COVER_OUT}")


# ========== DOCX ==========

def add_page_number(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run.font.name = BODY_FONT
    run.font.size = Pt(10)
    run.font.color.rgb = BROWN_RGB


def insert_page_break(paragraph):
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)


def set_page_break_before(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(OxmlElement("w:pageBreakBefore"))


# ---------- Content: the 50 tactics ----------

TACTICS = {
    "EMOTIONAL MANIPULATION": [
        ("Gaslighting",
         "\"That never happened. You're imagining things.\"",
         "Your specific memories of events keep getting questioned or denied.",
         "Document dates, screenshots, messages. Trust your records, not their retelling."),
        ("DARVO (Deny, Attack, Reverse Victim/Offender)",
         "\"How dare you accuse me of that. You're the one who hurt me.\"",
         "When you raise a concern, they immediately become the victim.",
         "\"We can talk about both. First let's finish the one I raised.\""),
        ("Love Bombing",
         "\"You're the most amazing person I've ever met. Nobody else gets me like you do.\"",
         "Intensity too fast. Idealization before they actually know you.",
         "Slow the pace. Watch what happens the first time you say no."),
        ("Guilt Tripping",
         "\"After everything I've done for you, you can't even do this one thing?\"",
         "Past favors weaponized to compel present compliance.",
         "\"I appreciate what you did. My answer is still no.\""),
        ("Silent Treatment",
         "(prolonged refusal to engage after a disagreement)",
         "Sudden, complete withdrawal of communication as a punishment.",
         "Don't chase. State your position once. Continue your life."),
        ("Emotional Blackmail",
         "\"If you leave me, I don't know what I'll do.\"",
         "Threats of self-harm, crisis, or catastrophe tied to your choices.",
         "Call a professional if the threat is genuine. Don't be the crisis line."),
        ("Playing Victim",
         "\"Everyone is always against me. You too, now?\"",
         "Chronic victimhood narrative that prevents any accountability.",
         "Engage with the specific issue, not the martyrdom frame."),
        ("Moving Goalposts",
         "\"Yes, you did X. But now I need Y. Then we'll be good.\"",
         "The requirement to satisfy them keeps shifting.",
         "Name the pattern. \"What would actually be enough?\""),
        ("Future Faking",
         "\"Once I get past this stretch, everything will change, I promise.\"",
         "Promises about future behavior used to justify current bad behavior.",
         "Evaluate based on what they do today, not what they promise about tomorrow."),
        ("Trauma Weaponization",
         "\"You know what I've been through. How can you bring this up?\"",
         "Personal history invoked to shut down current accountability.",
         "\"Your history deserves compassion. My concern still needs addressing.\""),
    ],
    "LANGUAGE TRICKS": [
        ("Straw-Manning",
         "\"So you think I should do nothing? Is that it?\"",
         "Your actual position replaced with an extreme version.",
         "\"That is not what I said. Let me restate it.\""),
        ("Whataboutism",
         "\"Well, what about the time you did X in 2021?\"",
         "Your concern deflected with an unrelated grievance.",
         "\"We can address that next. First let's finish this one.\""),
        ("Rapid-Fire Questioning",
         "(five questions in a row before you can answer one)",
         "You never get to finish a thought before the next question arrives.",
         "\"Let me answer one at a time. Starting with the first.\""),
        ("Semantic Manipulation",
         "\"I never said I would. I said I might.\"",
         "Retroactive reinterpretation of their prior words.",
         "\"Regardless of the wording, what I heard was X. Is that no longer true?\""),
        ("Hypothetical Trapping",
         "\"Imagine if I had actually done something terrible. You'd be mad about that too.\"",
         "Invented scenarios used to argue the real issue is trivial.",
         "Refuse the hypothetical. \"I am not here about what you did not do.\""),
        ("The Weaponized Compliment Sandwich",
         "\"You're so smart, which is why what you did was so disappointing, but I love you.\"",
         "Praise bookending a criticism to make the criticism uncontestable.",
         "Separate the parts. Address only the critique. Discard the packaging."),
        ("Deflection Through Humor",
         "(turns your serious point into a joke, or rolls their eyes)",
         "Real issues consistently met with jokes, sarcasm, or mock incredulity.",
         "\"I am not joking. I am going to ask once more, seriously.\""),
        ("Strategic Interruption",
         "(cuts you off mid-sentence, repeatedly)",
         "You cannot complete a thought. Always talked over.",
         "Silence. Wait them out. Restate your sentence verbatim when they stop."),
        ("Rhetorical Trap",
         "\"Are you really going to do this? Right now? Seriously?\"",
         "Fake incredulity used to pressure you into retreating.",
         "\"Yes. Really.\" And then nothing else."),
        ("Meaning Drift",
         "\"Well, loyalty means different things to different people.\"",
         "Words you both understood suddenly become negotiable when it benefits them.",
         "Pin down the definition in writing. \"For this conversation, X means Y. Agreed?\""),
    ],
    "SOCIAL PRESSURE": [
        ("False Consensus",
         "\"Everyone in the family agrees with me on this.\"",
         "Invoking unseen third parties who allegedly agree with their position.",
         "\"Name one. Specifically.\""),
        ("Fake Authority",
         "\"My therapist said you're the problem.\"",
         "Invoking an authority the other person can't verify or challenge.",
         "\"I'd like to understand their reasoning. Can I speak with them directly?\""),
        ("Public Shaming",
         "(raises your failings in front of others)",
         "Private grievances aired in public settings to pressure compliance.",
         "Do not engage publicly. \"We can talk about this privately.\""),
        ("Triangulation",
         "\"Your brother thinks you're being unreasonable too.\"",
         "Third parties inserted into what should be a two-person conflict.",
         "\"I'd like to hear that directly from him.\""),
        ("Social Proof Manipulation",
         "\"My last partner never had a problem with this.\"",
         "Comparison to unnamed, idealized alternatives.",
         "\"I'm not them. What matters is what's happening between us.\""),
        ("Splitting",
         "(plays mutual friends or family members against each other)",
         "People close to you turning against each other after interacting with this person.",
         "Talk to the other people directly. Bypass the splitter."),
        ("Isolation",
         "\"Your friends don't really like you. I'm the only one who gets you.\"",
         "Subtle or overt campaigns to reduce your outside relationships.",
         "Maintain your pre-existing relationships actively. Trust them."),
        ("Status Threat",
         "\"I don't know if this is going to work if you keep bringing this up.\"",
         "The relationship itself held hostage to your compliance.",
         "\"If this is a dealbreaker for you, let's discuss that directly.\""),
        ("Peer Pressure Framing",
         "\"Anyone in my position would do the same.\"",
         "Norm claims used to make their behavior seem inevitable.",
         "\"I'm not asking about everyone. I'm asking about you.\""),
        ("Appeal to Your Reputation",
         "\"You're too smart to fall for something like that, aren't you?\"",
         "Flattery used to make resistance feel stupid.",
         "\"Smart enough to notice what you just did there.\""),
    ],
    "INFORMATION MANIPULATION": [
        ("Selective Memory",
         "\"I never promised that. You're misremembering.\"",
         "Specific commitments retroactively forgotten.",
         "Get agreements in writing. Text messages are your friend."),
        ("Rewriting History",
         "\"Actually, you were the one who started that whole thing in 2021.\"",
         "Shared events reframed with you as the original aggressor.",
         "\"That's not how I remember it. Let's look at what I wrote at the time.\""),
        ("Information Withholding",
         "\"I didn't think you needed to know.\"",
         "Critical information reaches you after the moment it would have mattered.",
         "State the pattern. \"Finding out late is a trust issue, not a strategy.\""),
        ("False Urgency",
         "\"I need your answer right now. Otherwise...\"",
         "Artificial deadlines on decisions that don't actually need to be fast.",
         "\"I am not making this decision today. If that's a problem, I understand.\""),
        ("Premature Commitment",
         "\"Just say yes in principle. We'll work out the details later.\"",
         "Asked to commit before the actual terms are clear.",
         "\"I'll commit once I understand what I'm committing to.\""),
        ("Decision Fatigue Exploitation",
         "(brings up major decisions late at night or after a fight)",
         "Timing of big asks consistently when your defenses are lowest.",
         "\"Let's sleep on this and talk in the morning.\" Always."),
        ("Information Overload",
         "(overwhelms you with complexity, data, or jargon)",
         "You leave the conversation more confused than when you entered.",
         "\"Give me the three-sentence version.\""),
        ("Selective Truth",
         "(tells you only the parts that support their narrative)",
         "Details emerge slowly over time that change the picture.",
         "\"What am I not asking that I should be?\""),
        ("Document Destruction",
         "\"I deleted those messages.\"",
         "Relevant records conveniently disappear.",
         "Assume everything is preserved somewhere. Screenshot everything going forward."),
        ("False Dichotomy",
         "\"Either you trust me or you don't.\"",
         "Complex situations collapsed into binary ultimatums.",
         "\"Neither of those. Here's what's actually true...\""),
    ],
    "POWER PLAYS": [
        ("Intimidating Silence",
         "(says nothing, just stares)",
         "Prolonged silence used to unnerve you into filling the gap.",
         "Match the silence. Whoever speaks first loses."),
        ("Posturing",
         "\"I don't need this. I have three other options lined up.\"",
         "Claims of abundant alternatives used to pressure you.",
         "\"Great. I'll wait to hear from one of them.\""),
        ("Flipping the Script",
         "\"Actually, I am disappointed that you even brought this up.\"",
         "Any concern you raise results in you being the one on trial.",
         "\"We can talk about my behavior after we finish with yours.\""),
        ("Strategic Incompetence",
         "\"I'm just not good at that.\"",
         "Selective incompetence always in areas that conveniently benefit them.",
         "\"Figure it out, or we do it together once.\""),
        ("Feigned Helplessness",
         "\"I don't know what to do. Please help me.\"",
         "Helplessness used to extract resources, then disappears when resources arrive.",
         "\"What have you tried? What's your plan before you ask for help?\""),
        ("Limit Testing",
         "(pushes slightly past a stated boundary to see if you enforce it)",
         "Small violations used to probe what you will let slide.",
         "Enforce every time. The first small violation predicts the big one."),
        ("Slow Boundary Erosion",
         "(each ask slightly bigger than the last over months or years)",
         "The original agreement has shifted significantly without a renegotiation.",
         "Reset to the original in writing. \"Here's what we actually agreed to.\""),
        ("Malicious Foot-in-the-Door",
         "\"Just start with this small thing. It's no big deal.\"",
         "Escalating asks after the first small yes.",
         "Break the sequence. The second ask is where you say no."),
        ("Low-Ball Anchoring",
         "\"I was thinking $30K for this\" (for $90K of work)",
         "First offer is absurdly below any reasonable value.",
         "\"That is off by a factor of three. Let me know if you want to negotiate seriously.\""),
        ("Projection",
         "\"You're the controlling one.\"",
         "Being accused of the exact behavior the other person is engaged in.",
         "Note the projection silently. Do not defend. Their accusation is about them."),
    ],
}


def build_docx():
    doc = Document()

    # ==== SECTION 1: COVER (zero margins, full-bleed image) ====
    cover_section = doc.sections[0]
    cover_section.page_width = Inches(6)
    cover_section.page_height = Inches(9)
    cover_section.top_margin = Inches(0)
    cover_section.bottom_margin = Inches(0)
    cover_section.left_margin = Inches(0)
    cover_section.right_margin = Inches(0)
    cover_section.header_distance = Inches(0)
    cover_section.footer_distance = Inches(0)

    cover_p = doc.add_paragraph()
    cover_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cover_p.paragraph_format.space_before = Pt(0)
    cover_p.paragraph_format.space_after = Pt(0)
    cover_p.paragraph_format.line_spacing = 1.0
    run = cover_p.add_run()
    run.add_picture(str(COVER_OUT), width=Inches(6), height=Inches(9))

    # ==== SECTION 2: BODY ====
    body = doc.add_section(WD_SECTION_START.NEW_PAGE)
    body.page_width = Inches(6)
    body.page_height = Inches(9)
    body.top_margin = Inches(1.1)
    body.bottom_margin = Inches(0.9)
    body.left_margin = Inches(0.75)
    body.right_margin = Inches(0.75)
    body.header_distance = Inches(0.3)
    body.footer_distance = Inches(0.4)
    body.header.is_linked_to_previous = False
    body.footer.is_linked_to_previous = False

    # Header: logo right-aligned
    h = body.header.paragraphs[0]
    h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hrun = h.add_run()
    hrun.add_picture(str(LOGO), width=Inches(0.9))

    # Footer: page number centered
    f = body.footer.paragraphs[0]
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(f)

    # ==== INTRO PAGE ====
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run("HOW TO USE THIS GUIDE")
    r.font.name = TITLE_FONT
    r.font.size = Pt(20)
    r.bold = True
    r.font.color.rgb = DARK_RGB

    intro_paragraphs = [
        "The fifty tactics in this guide are the most common forms of manipulation you will encounter in adult life. At work, in relationships, from family members, from vendors, and occasionally from your own internal narrator.",
        "Each tactic is laid out in four lines: the name, the typical phrase or behavior you will hear, the red flag that tells you this is the tactic being used on you, and the counter-move that disarms it.",
        "This is a defensive tool. The goal is not to learn these tactics so you can deploy them. The goal is to recognize them in real time, which shrinks their power from whatever they were before down to almost zero.",
        "Screenshot any page. Keep it on your phone. When something feels wrong in a conversation, pull up this guide and find the row that matches. Most of the time, just naming what is happening to you in private is enough to change your behavior in the next thirty seconds, which changes the dynamic.",
        "The tactics are grouped into five categories. Skim all five on your first read. Come back to specific ones when you need them.",
    ]
    for text in intro_paragraphs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.4
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.first_line_indent = Inches(0.2)
        r = p.add_run(text)
        r.font.name = BODY_FONT
        r.font.size = Pt(11)
        r.font.color.rgb = DARK_RGB

    # ==== TACTIC PAGES ====
    tactic_number = 0
    for category, tactics in TACTICS.items():
        # Category header on new page
        cat_p = doc.add_paragraph()
        set_page_break_before(cat_p)
        cat_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cat_p.paragraph_format.space_before = Pt(18)
        cat_p.paragraph_format.space_after = Pt(6)
        r = cat_p.add_run(category)
        r.font.name = TITLE_FONT
        r.font.size = Pt(18)
        r.bold = True
        r.font.color.rgb = DARK_RGB

        rule_p = doc.add_paragraph()
        rule_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rule_p.paragraph_format.space_after = Pt(18)
        rr = rule_p.add_run("_" * 30)
        rr.font.name = BODY_FONT
        rr.font.size = Pt(10)
        rr.font.color.rgb = GOLD_RGB

        # Each tactic
        for name, what_they_say, red_flag, counter in tactics:
            tactic_number += 1

            # Name with number
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.keep_with_next = True
            r = p.add_run(f"{tactic_number:02d}. {name.upper()}")
            r.font.name = TITLE_FONT
            r.font.size = Pt(12)
            r.bold = True
            r.font.color.rgb = DARK_RGB

            # What they say
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.keep_with_next = True
            r1 = p.add_run("You'll hear:  ")
            r1.font.name = BODY_FONT
            r1.font.size = Pt(9)
            r1.bold = True
            r1.font.color.rgb = GRAY_RGB
            r2 = p.add_run(what_they_say)
            r2.font.name = BODY_FONT
            r2.font.size = Pt(10)
            r2.italic = True
            r2.font.color.rgb = DARK_RGB

            # Red flag
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.keep_with_next = True
            r1 = p.add_run("Red flag:  ")
            r1.font.name = BODY_FONT
            r1.font.size = Pt(9)
            r1.bold = True
            r1.font.color.rgb = RED_RGB
            r2 = p.add_run(red_flag)
            r2.font.name = BODY_FONT
            r2.font.size = Pt(10)
            r2.font.color.rgb = DARK_RGB

            # Counter
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.left_indent = Inches(0.2)
            r1 = p.add_run("Counter:  ")
            r1.font.name = BODY_FONT
            r1.font.size = Pt(9)
            r1.bold = True
            r1.font.color.rgb = GOLD_RGB
            r2 = p.add_run(counter)
            r2.font.name = BODY_FONT
            r2.font.size = Pt(10)
            r2.font.color.rgb = DARK_RGB

    # ==== CLOSING PAGE ====
    close_p = doc.add_paragraph()
    set_page_break_before(close_p)
    close_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    close_p.paragraph_format.space_before = Pt(60)
    close_p.paragraph_format.space_after = Pt(24)
    r = close_p.add_run("ONE LAST NOTE")
    r.font.name = TITLE_FONT
    r.font.size = Pt(20)
    r.bold = True
    r.font.color.rgb = DARK_RGB

    closing_paragraphs = [
        "The hardest part of defending against manipulation is not the tactical response. It is the act of naming, in your own head, that manipulation is what is happening to you.",
        "Most targets of manipulation spend years trying to convince themselves that what they are experiencing is something else. Their own sensitivity. A misunderstanding. A rough patch. Once you can name the tactic in real time, the power dynamic shifts before you have said a single word out loud.",
        "This guide is the naming tool. The counters matter, but the naming matters more.",
        "For the full system behind these counters, see Shadow Persuasion: The 47 Counterintuitive Conversation Tactics That Make People Say Yes Without Realizing Why.",
    ]
    for text in closing_paragraphs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.4
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.first_line_indent = Inches(0.2)
        r = p.add_run(text)
        r.font.name = BODY_FONT
        r.font.size = Pt(11)
        r.font.color.rgb = DARK_RGB

    DOCX_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(DOCX_OUT))
    print(f"Docx saved: {DOCX_OUT}")
    print(f"Size: {DOCX_OUT.stat().st_size} bytes")
    print(f"Tactics: {tactic_number}")


if __name__ == "__main__":
    build_cover()
    build_docx()
